"""
=============================================================================
dashboard.py — Dashboard Streamlit de Prévision Marine / Marine Forecast Dashboard
Point Sème (6.22°N, 2.63°E) — Golfe de Guinée, Bénin

METEO-BENIN / DPROM / SPAM
Auteur : LAOUROU MAKONDJOU DIANE
=============================================================================
"""

import streamlit as st
import hashlib
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import io
import sys
import os
from datetime import datetime, timedelta
import pytz

# ── Fuseau horaire Bénin UTC+1 ────────────────────────────────────────────────
TZ_BENIN   = pytz.timezone("Africa/Porto-Novo")

# =============================================================================
# AUTHENTIFICATION
# =============================================================================

def _hash(pwd):
    return hashlib.sha256(pwd.encode()).hexdigest()

def _get_users():
    try:
        users = {}
        for uname, udata in st.secrets.get("users", {}).items():
            users[uname] = {"name":udata["name"],"hash":udata["hash"],
                            "role":udata["role"],"points":list(udata["points"])}
        return users if users else _default_users()
    except Exception:
        return _default_users()

def _default_users():
    return {
        "diane":    {"name":"LAOUROU MAKONDJOU DIANE",
                     "hash":"841fccf7e520761479132298008eb9342d88c7650d2e9a96f593e47720df8d00",
                     "role":"admin","points":["seme","terminal"]},
        "wapco":    {"name":"Opérateur WAPCO",
                     "hash":"16eb9e9150beff2ff4c688a8c4655f416410c6e224b9e219a72de82e38b1db3c",
                     "role":"client","points":["seme"]},
        "terminal": {"name":"Opérateur Bénin Terminal",
                     "hash":"ee6cdf0e70179451a50180d326ba13ab6689e84d1b856fe0e21717b220fb44a2",
                     "role":"client","points":["terminal"]},
    }

def render_login():
    st.markdown(f"""
    <div style='display:flex;justify-content:center;margin-top:4rem;'>
    <div style='background:#1E2130;border:1px solid rgba(21,170,191,0.3);
                border-radius:16px;padding:2.5rem 3rem;min-width:380px;text-align:center;'>
        <img src="{_LOGO_URI}" style='height:80px;object-fit:contain;margin-bottom:0.8rem;'/>
        <h2 style='color:white;margin:0.3rem 0;'>METEO-BENIN</h2>
        <div style='color:#adb5bd;font-size:0.8rem;margin-bottom:1.5rem;'>
            DPROM / SPAM — Prévisions Marines
        </div>
    </div></div>""", unsafe_allow_html=True)
    _, col_f, _ = st.columns([1,2,1])
    with col_f:
        st.markdown("### 🔐 Connexion")
        username = st.text_input("Identifiant", placeholder="Votre identifiant", key="login_user")
        password = st.text_input("Mot de passe", type="password", key="login_pwd")
        if st.button("Se connecter", use_container_width=True, type="primary"):
            users = _get_users()
            user  = users.get(username.lower())
            if user and _hash(password) == user["hash"]:
                st.session_state["authenticated"] = True
                st.session_state["user_name"]     = user["name"]
                st.session_state["user_role"]     = user["role"]
                st.session_state["user_points"]   = user["points"]
                st.session_state["username"]      = username.lower()
                if len(user["points"]) == 1:
                    st.session_state["point"] = user["points"][0]
                st.rerun()
            else:
                st.error("❌ Identifiant ou mot de passe incorrect.")
        st.markdown("<div style='text-align:center;color:#4a6480;font-size:0.7rem;margin-top:1.5rem;'>© 2026 METEO-BENIN</div>",
                    unsafe_allow_html=True)
UTC_OFFSET = timedelta(hours=1)

# Logo METEO-BENIN
_LOGO_URI = "data:image/png;base64,/9j/4AAQSkZJRgABAQAAAQABAAD/4gHYSUNDX1BST0ZJTEUAAQEAAAHIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAACRyWFlaAAABFAAAABRnWFlaAAABKAAAABRiWFlaAAABPAAAABR3dHB0AAABUAAAABRyVFJDAAABZAAAAChnVFJDAAABZAAAAChiVFJDAAABZAAAAChjcHJ0AAABjAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAAgAAAAcAHMAUgBHAEJYWVogAAAAAAAAb6IAADj1AAADkFhZWiAAAAAAAABimQAAt4UAABjaWFlaIAAAAAAAACSgAAAPhAAAts9YWVogAAAAAAAA9tYAAQAAAADTLXBhcmEAAAAAAAQAAAACZmYAAPKnAAANWQAAE9AAAApbAAAAAAAAAABtbHVjAAAAAAAAAAEAAAAMZW5VUwAAACAAAAAcAEcAbwBvAGcAbABlACAASQBuAGMALgAgADIAMAAxADb/2wBDAAUDBAQEAwUEBAQFBQUGBwwIBwcHBw8LCwkMEQ8SEhEPERETFhwXExQaFRERGCEYGh0dHx8fExciJCIeJBweHx7/2wBDAQUFBQcGBw4ICA4eFBEUHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh7/wAARCABIAGcDASIAAhEBAxEB/8QAHAAAAgIDAQEAAAAAAAAAAAAAAAcFBgECBAMI/8QAPRAAAQMEAAQEBAMFBQkAAAAAAQIDBAAFBhEHEiExE0FRYRQiMnEWgZEVQlJisQgXI0OiJCVTcoKDobKz/8QAGwEBAAIDAQEAAAAAAAAAAAAAAAIDAQQGBQf/xAAjEQACAgEFAAEFAAAAAAAAAAAAAQIDEQQFEiExQQYTFVFx/9oADAMBAAIRAxEAPwD7LoorBIHc96AzRVFyriXZbTdVWC1sTclyIDrbLSjxVt+hdX9DI91kfaopcTi3kKFO3G+WnB4Oifh7eyJ0vl/mec02k/8AKk/c0AzFLISVa6ffVRs7JLBbz/vC+WqH7PzG0f1NLuRwxxN1xf4huGR5NLDXip/al3eKVjzAbQUo17arrgYNwvghxcXBscIDSHEqMFC+YKOu6gfPVaktbRGWHImoNlwRmmHOKCG8ssK1HsE3Fon/ANqlYk6JLG4sqPIHq06Ff0qrtYLw7fW4z+CsaJSAF6tjPQny+nvUVP4P8LHCVpxOBCcJ0HYK1xV79lNKSa2IzUllEBihW1aIIrNLBvh1dLYsHDOI+S2so+mJcFpuMbXpyu/OB9l1lzLeIGKA/jPFEXi3p+u7Y4VOlsfxORVfOB68hVUwM6iobE8lsWUWpN1sF3i3KGrp4jK98p80qHdKvYgGpmgCiiigMLPKknp09aU0y8XvildH7Zitxfs2GxXSxOvkfpIuS0khbMRX7qARpT3mdhPrUjxhkzrrKtHDq0ynY0rI3F/HPtq0uPbm9GQoEdlL2lsH+c+lXyyWyDaLVGtduhtRIURpLMdhsaS2hI0Ej8qA4MVxWxYrakWzH7cxAiDqtDafmdX5rWs/MtR8yokmvHILqzD0j4tplwb228hRS4n2IqwOa5OtUjLpT6VLbYmznCglSmkBLSGx5lTyh0A/WtDcZuFLa9LakpSSZBu3EvPpZYWlCN86T4pW2gjqFIV3A7gpNbSVKZjrV8WhwaKfDA0SFL5ike5OgPSqDKu70e6SFFfiIWByq5ubm996G/vXLcrzKlxwyytaFhQWlaT1Tr2r5DLXXz1+LHiGe2df+Jao5V9vA0Yl8CCpp2V4S0rU48Ak/M4ry9dJH6muZF1uUuQW2JDTqGBypdUoISQfv50sYVoya7qD0PmkknS9uBJSR67PT2rF7YzPG3ExJ0J5IcTzpWghSFb/AJh0r7Htq08qE6JZWPk4vUq2qxxmsMZsPKlwJi/i3CJKNpKD1GvX335Vc8Ov7OQQXXG1APMr5HNH9CPY0hLfj2bX1gT49uecaUQgKUQgE+29bHvTl4S4rNxqzvqubyFz5jiXHUtnaGgBoIB8/PZreko46ZVByz2cuZ8OGX7mrKcMmJxnKgOstlH+zzdf5cpodHEn+L6hvYNd3DjNl5EJdnvFu/Y+UWopRc7apfME7+l5pX+YyvqUq/I9RVzKQdbAOu1LfjVaZcOPG4h2BlX7cxkF5SUdDNg95EZXqCkFafRSRrvVZaMmiuSzTo1ztUW5QnQ7ElsofYWP3kLSFJP6GigKFbB8R/aQvTj3UwcYitMb8g7IdUvX5tppkUss2dGLcY8ay54hNsu8dWPTnD0Sy6pfiRlKPoVhSNnzWPWmYlQUNigBY2NVWcygB+KC8ZEptPzNwkJHK84PpCz/AA70dHpVnJAGzXi8Ukbqq6tWQcWZjLi8iRvFmWLkpEpsS30uIjuuOdS46UeIsgj6UpHRKR3rlbaYTH21CZhoUhpZfUSoNpcB5Fq/l2AD6bpxO2aIp+KUJCW2HlSOQDZW4oEcxJ+5rit2MtRRBXpBMfxUOAjoppZJ8PR7pBI6H0rkNT9K03282j2K93tjBRKXj1iWZKWnIMRUhtHifCzo+yU76+E8k6WnfUbB10r1M/KIfFiNbTbW4uHx4rinFLHyIJTsOLWrp0OwEjt3piWy3x4MRiOkFSI4UGif3Uk9h7eVc+U2WJkWPyrPM5gxJRyL0ffp9+tdVt2lhpYcMdeHi66dl75J9p5/ovstZyXNbhjF6w27tOWBmUl5YSvlbcKHCCpXmpOgdD1pssjY3uq5w/xiHh2OM2GE846y0pa+ZzuVKOz0HRI69hVkSpAGugrdsmnhLxGvRW45lL1m9ecphuTGdjvJCm3UFCwfMEaIr0qn8XMk/C2C3O5tBTk9xAiW6Ok6U/Kd+RlA9SVEHp2AJ8qgbJG/2cHFucE8aStRPgsOMIJ/gbdWhH+lIoqw8Nsf/CuB2THCoLXb4TbLix++4E/Or81En86KA6cwsFryTHJ1lvLCpEGY14bqAdKHXYUk+SknRB8iBVIwnLbnjt3Z4f588kXXXLaLwscrN5aT2+bsmSBrmQe5+ZOwejPqHy7GrLlVndtF/trNwguaJacH0qHZSSOqVDyUNEUBKr6t9qV+Z5jllpu2SRrTaWJseFEQuO+pWhGcLRWoujupGh01130rZu38R8BRyWt1eeY+39EWU8lq6Rk+iXT8j4A8l8qvc1IWPifgdymOW6fLFhuro5Xrfe45hvnQ7f4mkr/6SRWU8ekZJtdPBXk5bkz0TJXWrzFjP2tKlxm1R0L50oSkq2Obm7q67Gh01W1yzPJbNcbhGkXGJM+DtrLwCo6Gw444N85AVzcqR8xCQRpJ2RTRbhW1SlyGYcVSnk6W4lpJ8Qe5HcVlcC3qk/ELhxy8E8odLQKgn03317VPlH9FX2p49KVY8gunxN7iuXOHemoVvTLanx2wlIWUq22oJJSfpChrro9aq+OZ9f7nIsbEi4xGPiFO/FrDSBzBPhcoG1cvZZHybPbp3pvRYEGKwWIkOPHZJJLbTQSkk9+g6da0/ZdsCWkC2RQljq0nwE/4Z9U9On5UUo/KEq5vxioj8RL+LdHYmNR49y+EXK34ZLchrxW0IcR6aClBSfJQ96m8ZzK7SMog2K6sNtOPzZaWHko+SSw3zAaPZLiSAFD0IPnVvyK6YtYIfxd+nWm2MNg6XLcQ2AO5A5v6CqSriIi/rS1w1xGVkq0qPJcXmzCt7RPdXjLG1/8AbSd+tY5LHgjXJPtjByW+WrHLNIvN7uEe3wIyeZ599WkpHp7k+QHUnoKX+JQrnnuWRM+yGE/b7PB5jjdpkJ5XNqGjNfT5OKSdISfpSST1Nddm4cyrhdo+R8Q7o3kd3jq8SJES2UW63q9WWj9Sx/xF7PpqmKgEJHN3qBeCEBO9edFbUUAUUUUAEbHnUffLHZr7DVDvdqhXKOru1KYS6n9FA0UUBSFcGcOir8TH3b9jS99rPdXmED38PmKP9NaJ4d5VHHLbeLuWtI7alMxZJ191t7/OiigM/gTPfPjJfyPPVqhA/wDzo/uukyk6vXEfOrgk9FNouCYyFD3DSEn/AM0UUBJWLhNw+s8pM1nGosucnr8XcFKlvE+vO6VHfvVzDSAEgbCU9kjoBRRQHpRRRQBRRRQH/9k="
_LOGO_BT_URI = "data:image/png;base64,/9j/4AAQSkZJRgABAQAAAQABAAD/4gHYSUNDX1BST0ZJTEUAAQEAAAHIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAACRyWFlaAAABFAAAABRnWFlaAAABKAAAABRiWFlaAAABPAAAABR3dHB0AAABUAAAABRyVFJDAAABZAAAAChnVFJDAAABZAAAAChiVFJDAAABZAAAAChjcHJ0AAABjAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAAgAAAAcAHMAUgBHAEJYWVogAAAAAAAAb6IAADj1AAADkFhZWiAAAAAAAABimQAAt4UAABjaWFlaIAAAAAAAACSgAAAPhAAAts9YWVogAAAAAAAA9tYAAQAAAADTLXBhcmEAAAAAAAQAAAACZmYAAPKnAAANWQAAE9AAAApbAAAAAAAAAABtbHVjAAAAAAAAAAEAAAAMZW5VUwAAACAAAAAcAEcAbwBvAGcAbABlACAASQBuAGMALgAgADIAMAAxADb/2wBDAAUDBAQEAwUEBAQFBQUGBwwIBwcHBw8LCwkMEQ8SEhEPERETFhwXExQaFRERGCEYGh0dHx8fExciJCIeJBweHx7/2wBDAQUFBQcGBw4ICA4eFBEUHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh7/wAARCABoAUQDASIAAhEBAxEB/8QAHAABAAICAwEAAAAAAAAAAAAAAAUGAgcBAwQI/8QAURAAAQMDAQUDCAUHCAUNAAAAAQACAwQFEQYHEiExQRNR0RQVFiJUYXGUCDKBkfAjRKGxssHhFyQzNFJVdJMlQmSDkhgoNkNFVldydYKEw9L/xAAbAQEAAgMBAQAAAAAAAAAAAAAAAQQCAwUGB//EADMRAAEEAQIDBgUDBAMAAAAAAAEAAgMRBCExBRJREyJBYaHhBjJxgZEUQlIVscHwFiPx/9oADAMBAAIRAxEAPwD7LRFwThEXKLHf48RhC8ZAyPvUWosLJFwHAoCpUrlEREREREREREREREREREREREREREREREREREREREREREREREREREREREREREREREREREREREXDvqla121a+k0jbWUdvaHXWrjc6ne76sQGAXHPPnyVs11qKm0vpqqvFSC4RM9RjTxe48gF8p6svt41NefOF5qBNOwb0AaMMZEebQP1nrhdjhHDjlP53DuBeZ+IeMDEj7KM98+ik9RbSdZX2jZTVd17CJrQT5I3snP449Yjj9yjNM32ptuqbdc66suNVBRVHbOj8pcS4NHAYccdVHtpjjd4HAc3j1HBw/QktMWl3qkkO/ef4fcvYfo4GsLGtABXz13EMp0gkc4khfUWmdpml77Q1FVBVOp5IQ9zqaowyUhoySG54ro05ta0Vemki6Nt8m9jsqz8m4+8dMfavl2aBxOSACBzXjlYBlvMcjwXI/43AQe+fJegb8YZWlsHmvt2yXm23qibW2qshrKZzi0SxOy0kHBCkAV86fRl1Q+ku02lquUNgqsz0ucNDZB9Zo+I4/YvopnJeazsU4sxjK9vwvPbnY4mCyRE6KmuiiKgbXNezaFitkkNrNykr5jC2MTbmCBw6cckqvDaXtA3sfyV3P8Azv4KzHiSyNDht9QqMvEIYnlhux0BK3AuCVpx+167W65UMeq9E3CyUdXMIm1LpN4Nce8Y5cVMbR9o1fpTUttsdt09Jeam4QGSNjJ9xxIOMAY7uKk4U3MG1v5jwUf1LHLC69q8Deuyul81JZbG+nju9xp6R9S/s4GyOw6R3c0cypZj94AjrxXzxTVusXa4l1fdtmt5udbu7lFFJI0R0TO5gxjePVx4rYFh1xqaoorzV3nRVXZoLdROqY3TTAiUjPqD34HNZS4RjaKN/cLVBxISOIcCPsVslFqXZjtntmqbmy2XKk8z1U7Q6kEk28yf3A4HH3dVtfe5jPBaMiCTHdyyCirmPlRZDOeM2FmioVo10+4bSr5pE2/s/NcAmFR2me0zjhjHDmoey7VX12zy+atNo3PNdQ6DsO2/pCCBnOOHNZ/pZj+3p67LU7PgBNnr6braqLWeybazbNazvt1TTm2XRo346eSTeEzP7TDgZ+C2UXLCWF8LuV4ordj5EeQztIzYWSKi6A12dV6g1FahQGlFmqRDvmTe7XmM8uHJXcOOCsHsLDR3WUczJW8zTos0Wp6/aLrynuFRDFsuuk0McrmMlE4w9oOA7l15qLt+2TU9yrKyit2zysqqiifuVMcdRl0Tu48PcrLMGZ4sV+QqT+K47Dyuu/oVuxFq+g2hap8zXm6XnQ1XaobdRuqI3TTjE7hzZnHDvUNa9rWsrrb4bhbdm1wqqSZu9FLFUAteO8cEGFKbqtPMKXcUgbW9nwordKLUUe0nXr5GMfsruoa5wBPbcgTz5Lv1htQu1l1q7TFo0lNeKoU7ZwIpw1xBGTwx0UDClLuQUTvuFkeJQBnObH2K2suMrUP8pW0P/wAKrn/nfwXu0btTkuWqmaZ1BputsFzmYXwsmcHNkxx/Vn7lLsKVovQ/QgqG8SgcQDYvqCFtFFpyj220fp/Ppm52w0VLHVvpRXGfLQ8HDcjHAE/ctvRuyBxytU0D4a5xVrfj5cWQCYzdbrtRQ2rb9RacsdVd7hL2dPTxlzj1cejR7yeAVA2ZbVpNaS3hjrGaAW6n7dpM++ZBx4HgMckjx5ZWGRo0CxlzYYpBG46lbYRanoNq5qdklVrjzNumCp7EUpn+t6wGd7HvUhsp2o2zWnaUb4DbrrFxdSyPyXs/tNJ5j9SzdhzNY5/LoDRWA4ljl7Wc2rhYWyEXWX4aTjkqZsz1w7WU97j8hNJ5srPJv6Xf7Tn63u5clpDCWlw2Csvlaxwadyrui6XOOeCLCwtlrWP0khTO0PTPmlMUzK1nYH/VLsHId7sZWg6eEsLgAGFg390nO47vHew8j3ZX1brnTNNqexS22pkdE/60MoGezeOTsdfgvmittVfY7g633KmfTSwvO5lpDXccbzfcfcfsXq+BZTBEYr1u18/+KcOU5LZq7tUumCma5uN0+rwAPQceH2ZI+wLmeBu9nGPcpGigklLhBDLUEN+rGwk/d+5dODKC4xPjwcYeMH7ui7YlBNArzvZnl2UJVQDDz3qLqGEEbuOPeAP0lWKsY0ccclC1zG74DDnHE4LCf3lWGuVGRoGq8turpLZdKO5wNa+WjnbMwE8y05xn38l9k6N1Fb9TWGmutvqIpWyxh0jGuyY3Y4tcOhBXxfMXNOXOeD0yFO6D1xetE1tVVWlkM/lMe7JDOTuF3R+B1XL4tw05jQ5nzBd/4e4x+gkLH/I70X2bvA8MrnK1ZsEvurtVW6uv+oaiJ1FO9sdFHHEGAbud9w64zw+xbRAK8TPEYZDG7cL6TjzieMSAUCtI/SpMw9EjTsa+YXEmJrjgOeN3dB+JwpNt7247oB0hYMD/AGvj+0vZt30fqDVTLFJp9tM6ot9UZ3dvJujpu/HiF5TNt0IOKDSw/wB4T+9dON7DjsHdsXuVxZopBkyO7wBr5VQtr1Xr6uorZUa7sVLQWKlrY3yuoZg95cTjqe7Kmdr9bcDtZ0VWadp4qqtfRl1HHK7dY/OcBx6cCvXqbS21jXFHDZdSS2Ght5mbJI+lJLzu9w4/uUptJ0VqyfVmnL5pOGhmdZqTsmeVzbo3hwGR14KxHkRBzA7lFB222oVV+LM5ry0ONlu++iyF625lmBpDT/A+1n/9KQirtd1mi9TN1nZrfbQ23yCn8ll39/1Hb2eJx0XiE+3ccPN+lh/vD4qQt1PtNudpvdBqilssbKi3yRUhpZDkyuBGHZ6Km+hR7v2KvsD6IJedDuNFRdn+gLfrjYbaQ9/ktzp3zGjrG/WjO+eB6lp/RzCsGzLaBcbVeRoTaB/NbtERHS1cn1KgdAXcsno7r8VbtjOnbhpbZ9RWW6CHyqF7y/s3bzeLs812bTdCWvW9mNLVfkKyIZpatg9aJ37294UyZLJpHRy/KTofEfTy6qIcOSGFk0PzgajwP1VE0ef+cdrQEcHUA4/AMVO0i0f8njXH/qMn7TFdtkug9Y2HWlyvOpKimqu3oPJmzsm3nSOBGCRjPIdV0ae2c6motkOptMzspBcbjVvlgaJgW7pLSMnpwBVp80bXUHXXL6Km2GZ7Q4sIPf8AXZYRbOYtU7MNNXm1T+QagpKCJ1NUM9XfwM7riP0HorBso2kSXKrdpPWDPN+pqUljmvG6KkDqPfjj7+YVx2eWyqsuibRaK4M8ppaVkUu47ebvAccHqq9ta2c0er6NtdRSeQ36kG9R1jeBBBzuuPdnkeipGdkrjHMdL0PT2V4YskDRLANaFjr7qrfR/DRtA2hYP/aI/aet0har2DaL1Ppisv8AW6mNM6e5SRyCSKXfL3DO8Tw4c1tXdwtWe9r8glpseyucKY9mOA8UdT6rhx4c+q0vsLydq+0X/HD9t63S4cFrfZdpG8af15rC8XFkIpbtUiWlLJN4kbzjxHTmEx3tbFKCdwK/KjMjL8iFwGgJ/sprbHw2X6j/AMBJ+pas2XXPavTaDtUOntM2estjYj2E01TuvcN45yMjqtxbRLXVXzRN3tNCGGpq6V8UYe7Dd4jqei1bpK1ba9NWCkslvo9OvpqVpbGZZSXYJJ44+Ks4j2/pyw8t3eqp58bzkteOaqruqborztpdXU7avSNjip3StEz2VeXNZkbxAzxOFUtaXessP0ipLjb7NU3ieO2tY2mpyN9wLeJ+xWtlRt2dI0SUGlw3ILiHnOOvVeyn0de3bdDq+VlN5t83iDeEnr9pu4PqrOKSOJ7nODa5ToPErVLDNKxrWFxPMNTWijBtX1V12Vag/wCP+CrFju1brT6QNprrpbX6fnttK4xUdST2sww7lwH9on4BfQu57z961xqLSF4rNtdi1XTRwebqKkdFO50mJMne5DrzWqDIhHMAwNNHWyt+Vhznkt5cLGleq1xpnSVJrS/bTLRUt3JvLw+lmIz2MoL8H4d6vOw3WFbU0lVo/U35K/2b8m8vd/TRt4B2epHD4ggr3bL9I3rT+sdXXW5xUzaa7VQlpjHJvO3ck+sOnNR22XZxc77X0+o9ITx0V7jYYJiZDGJo3DGcjqMke8fYtks0c0hje7TSj0NKvDjTY7O3YO9rY6i1Xr/PPte2lR6eoZT6KWaQSV0reU7wfqg9eOQPdkrLZ1DFT7UNpUELAyOKnLGNAwA0A4AWzNmWj6PRel6a1U5D5sb9TNjBlkPMn3dyrOkNF3y3631xd6psHk15a4UhEmXdfrDopbkx0+NujQAB566lH4UvNHI4W9xJPlpstXWUgfRQupweFw/+xqvV62c+kOjdP6n03KLbqaloIHxzsO722GDDXHv9/wBhyF5KHZtqqn2C1ujjHSedpqsytHbZj3d8H63wC23pKhqLbpe126q3e3pqSOKXdORvNaAcHuU5OXyW+N194n7aKMPCMh5JW13QPodVTNlO0hl/dLp2/Qm3akpCWSwSDdE2ObmZ694+0cFB/RlOajWo5Ft3d+9WLats3h1QyK7Wmo826go8Opqph3Q8jkHkfoPRVvZ7o3XultH6nibFRG+XSZr6ZwqvVaS3Dnk44EE5wsCcd0Dyw0XVp01/stnLksyGB7bDL166aD6qK2obbrjZNYVVo09TUdXTUgEcssmeMozvAceQ4D4gorBo7YjpyHT9P6VUcdyvD8vqpu0cBvE/VGDxx39eKLa2ThzAGlpNeK0Px+JvcXAgX4LcDhvLxXC0W24PY+uoKapfH9R0sYcW/DK94RcQGtl6dzQ4UV5qWgo6XPk9LDFnnuMAyqHtntVBHpWouMduidV9qzMzWes0Z4knuwtirzXGlhraKakqGB8UrC14PUFb4JjHIH2quVjNlhdGBVhfI9XMWxuIkw1vN3HAz1ODkD38QoWvdjIcw4dyLyHh3vDvFTGq4zar9X29mY2UtS+IHGHRYPD/ANhHRV2qee0duxsa48XMH1T7wF9EgPOwP6r45k9x5YdwvO9wDcAmPP2hXrZHs0qda1Rq7kZqayQu9aWMYdUHq1hPIDq5UCodvMcGDGRgtJ5FfXmx2soK3ZtY5rd2bYW0rYy1gwGubwcP+IFc7jeW/HgAZu5dz4a4fHl5H/Zs3VWiz26ktVtp7dQwthpadgZFG0cGtC9oxhYtII4HK56Lwh1NndfUGgAUNlHXu601rZG+ojqXtecDsYXSH7cDgoz0ttfs10P/AMGTwU/NIyJjpHvbG0Di5xwB9q8Qu9q/vSi+Yb4qpM4td84C2N5SNrUb6W2ofm10+Rk8E9LbWePkt0x/gZPBTpqITH2omb2eM7+96uO/K8wu9r/vKk/z2+K18z26dqPwsg1lXyqL9LbX7PdPkZPBPS21+z3T5GTwUyK2jMRmFXCYgcF4kG6PtWHnS2/3lS/5zfFTzP8AGUf791IDfBqiW6ttXs10+Rk8Edq21+z3T5GTwUzDW0cocYqyGQMGXFsgIb8ccl1G72oHBulGCOf5dvioLneMo/37qRW4aVFjVlr9munyMngnpba/Z7p8jJ4Kbgq6WeMyQVUMrG83MkBA+5ZQTRVDO0gmZKzON5jshG9odpB+PdY90eCghq21+zXT5GTwXPpdbMf1a6fIyeCnJJo43tY+VjS7g0OeBvHuC66ivoqZ27U1kEDj0kkDf1qS143kH4U9w/tUO3V9sH5tdPkZPBZemFs9mufyMngpOK522WQRRXGkke7k1szST+legSxmXse0b2gGdze4478I0vI0lH4QlvRQZ1ha8caa5/IyeCxGr7WT/Vrp8jJ4KwPc2NhfI8NaOJJOAEhc2Rgex4c1wyCDkFZck1/OL+ijueLVAHV9rx/Vrp8jJ4INXWv2a6fIyeCsW7+Mpu+/9Ky7Oc/vH4UAsGwVd9L7Yfza6fIyeC5GrrWPze6fIyeCsIbx/iud38ZTs5v5+nunc6KvemFs9mufyMngsTq61nnTXT5GTwVk3Vxu/jKgxTnd4/HugLOirnpfa8Y8munyMngnpda8H+b3T5GTwVix+MrkN5+KdlN/P091JLOirnpda/Zrp8jJ4J6W2v2a6fIyeCsWPxlc4U9nP/P091Ft6KuHVtr9munyMngnpbax+bXT5GTwVj3fxlN38ZUCKYGw/wBPdTzM6eqrnpbaz+bXT5GTwXHpZa/Z7p8jJ4Kybv4yud38ZU9nP/P091Fs6eqrfpba/Zrp8jJ4IrJhFPZz/wA/T3S2dFyiIrSwQrorHPZTSOjbvPDCWjvOF3rFzchSFi4WF8TX2vq6q51FZcJHyVL5XCYv5n1jz+HJRcryPyZOccWO93crtt3tcdn2kV8cUQjgqw2oaByO8PWP3grX5fn1c5A5L6TgvEsDXt2IXxnMxnRZL2O3sqSstoumorrDbbNSOqayXk0cAB3uPQe9fQlHVz7HtlVBS10EdfcJJ3js2P3Y2vflxG8RyH6VSfov2O6y6ll1HE+NttiikpZhves95wQMe7nlXbb9prVmqa+x0VjYZaDfcKn1w0ROJAD3Z5gNzjC4HE8lkuYMeQ9waler4ViOxsE5MQPaHQKw7DK66XTRjr1dZzLNcayWoYCMbjCd0NHuGOC2ACvBY7bBabTSWymaGw00TYmY7gMZXvAXl5nB8jnNFC17TFY6OJrXGyAqVtKjFbctO2uYl1JVVpE8YOA8BucHvCkzovSuMeYqMj/yLu1Zp5t9ipnMq5KOqpJRLTzsGSx3wPMKN8yay/73xfIN8VxJIiJnufHzg1W3+V02uBYAHVSrepLXDadPQ2ilr3TUdTeGMdG12OzjJ4xc+QVyGj9LBo/0HRYxw/JqNZoaE2CW3T3GaSqlqhWPrAwB3ajkQOXRdjbLrJrd0auiwOWaBuf1qtDjujfb4bBA6aV9Vm6UOAAdR8fNRNrs1rj2g3exsoohbJaCKV9P/qb+9zx0XRqfTVgpdXaepKe00zIamSXt2BvB+G8M/Aq0ad07UW+5VN1uFydcK+pY2N8nZhjWsHIABd93sTq+/wBqunlO4Le57uz3M7+8Mc+iHh7nwFrmD5vwLRuQGyWD4f4VY19ZrbadLGO00cdF5ZWQQzmL1S9hfyyrDHonSoYP9B0ZIHVmV7NT2SC/WaS3VD3xbzg9kjObHNOQR9qhmWPWLGho1kxwAxl1A0n9asOx+znLxHzChW2n5WDZLb89H7qErKKjsusLjQWuFtLTT2SSWWKP6peCQDjoeK8Gl4bnpnTtu1FaxLV0E8DXXCjySQT/ANYz394VrodJVbZ66tuV5krq+qpTSibsgwRsOeTR7yprTlr802CktTpBP5PEIy/dxv464VOHh0jpS8jlq68iStxyGhvKNdr81U9SXOhulfpGvoZmTQSXDIIPXdPA+8Lr0dZrVf6y+1t4oYq2dtykia6bLt1jeQHcF7qnQVN6UU14oao0sUc3bvpQ3LHPxjeA6Ervdpa60lfVz2O/GgiqpTLJC+nEgDzzIJ5ZUDGye2MkrOYXt9qsJzxBnKx1f+2o7aBpiw2/SNfW262U9JV08faQyxNw9jgRggrOxSOm2lRyvcC59iic455kuXpr9KX+6QmjuuqnTUTyO1iipGsLxnON5eq66XqpLpBdLPdDbauKn8mceyEjHxg5AIPULN+NKZRKyOgK00vS1iJGCPlLrJvVSesyPRS6j/ZJP2SurZ//ANCrP/g4/wBlRNXpzVNbTSUlXq0Op5W7koZRNDi08wDngrPZ6GO226noIM9lTxtjZnngDC6MPayZJkc2hVa9bWlxa2PlBs2vYuQuMLldEKuiIilEREREREREREREREREREREREREREREQoiItPfSa0y256Obe4IZH1lukBHZN3i5jjhwIHEjkVofT2gdY3u5QUdLYqyFk2P5xPEWRsb/AGiTz+AX2s5oPAgEdyYXWxOLy40XZNGi4GZ8PwZU/bONFV7QWmKPSWmKSzUbW4ibvSvA4yyH6zj7yVYNzjlZgLnC5b3l5LjuV24omxMDGjQLEDC5XOEWAWa4XCywiUlLHATAWSKUpY4TCyRRSLHCYHcskUpSxwmFkiilKxwEwskU0oWOEwskSkpY4XIXOEUUpRERERERSiIiIiIiIiIiIiIiIiIiIiIiIiIiIiIiIiJhERERERERERERERERERERERERERERERERERERERERERERERERERERERERERERERERERERERERERF//9k="

def now_local() -> datetime:
    return datetime.now(tz=TZ_BENIN).replace(tzinfo=None)

# ── URL CSV GitHub ────────────────────────────────────────────────────────────
GITHUB_RAW_URL  = "https://raw.githubusercontent.com/DianeLaourou/Ecmwf_open_data-point-based-forecast/main/"
GITHUB_TREE_URL = "https://api.github.com/repos/DianeLaourou/Ecmwf_open_data-point-based-forecast/git/trees/main?recursive=0"
GITHUB_OWNER    = "DianeLaourou"
GITHUB_REPO_BT  = "Benin-Terminal-Forecast"

# ─────────────────────────────────────────────────────────────────────────────
# CONFIG PAGE
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Marine Forecast — Sème | METEO-BENIN",
    page_icon="data:image/png;base64,/9j/4AAQSkZJRgABAQAAAQABAAD/4gHYSUNDX1BST0ZJTEUAAQEAAAHIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAACRyWFlaAAABFAAAABRnWFlaAAABKAAAABRiWFlaAAABPAAAABR3dHB0AAABUAAAABRyVFJDAAABZAAAAChnVFJDAAABZAAAAChiVFJDAAABZAAAAChjcHJ0AAABjAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAAgAAAAcAHMAUgBHAEJYWVogAAAAAAAAb6IAADj1AAADkFhZWiAAAAAAAABimQAAt4UAABjaWFlaIAAAAAAAACSgAAAPhAAAts9YWVogAAAAAAAA9tYAAQAAAADTLXBhcmEAAAAAAAQAAAACZmYAAPKnAAANWQAAE9AAAApbAAAAAAAAAABtbHVjAAAAAAAAAAEAAAAMZW5VUwAAACAAAAAcAEcAbwBvAGcAbABlACAASQBuAGMALgAgADIAMAAxADb/2wBDAAUDBAQEAwUEBAQFBQUGBwwIBwcHBw8LCwkMEQ8SEhEPERETFhwXExQaFRERGCEYGh0dHx8fExciJCIeJBweHx7/2wBDAQUFBQcGBw4ICA4eFBEUHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh4eHh7/wAARCABIAGcDASIAAhEBAxEB/8QAHAAAAgIDAQEAAAAAAAAAAAAAAAcFBgECBAMI/8QAPRAAAQMEAAQEBAMFBQkAAAAAAQIDBAAFBhEHEiExE0FRYRQiMnEWgZEVQlJisQgXI0OiJCVTcoKDobKz/8QAGwEBAAIDAQEAAAAAAAAAAAAAAAIDAQQGBQf/xAAjEQACAgEFAAEFAAAAAAAAAAAAAQIDEQQFEiExQQYTFVFx/9oADAMBAAIRAxEAPwD7LoorBIHc96AzRVFyriXZbTdVWC1sTclyIDrbLSjxVt+hdX9DI91kfaopcTi3kKFO3G+WnB4Oifh7eyJ0vl/mec02k/8AKk/c0AzFLISVa6ffVRs7JLBbz/vC+WqH7PzG0f1NLuRwxxN1xf4huGR5NLDXip/al3eKVjzAbQUo17arrgYNwvghxcXBscIDSHEqMFC+YKOu6gfPVaktbRGWHImoNlwRmmHOKCG8ssK1HsE3Fon/ANqlYk6JLG4sqPIHq06Ff0qrtYLw7fW4z+CsaJSAF6tjPQny+nvUVP4P8LHCVpxOBCcJ0HYK1xV79lNKSa2IzUllEBihW1aIIrNLBvh1dLYsHDOI+S2so+mJcFpuMbXpyu/OB9l1lzLeIGKA/jPFEXi3p+u7Y4VOlsfxORVfOB68hVUwM6iobE8lsWUWpN1sF3i3KGrp4jK98p80qHdKvYgGpmgCiiigMLPKknp09aU0y8XvildH7Zitxfs2GxXSxOvkfpIuS0khbMRX7qARpT3mdhPrUjxhkzrrKtHDq0ynY0rI3F/HPtq0uPbm9GQoEdlL2lsH+c+lXyyWyDaLVGtduhtRIURpLMdhsaS2hI0Ej8qA4MVxWxYrakWzH7cxAiDqtDafmdX5rWs/MtR8yokmvHILqzD0j4tplwb228hRS4n2IqwOa5OtUjLpT6VLbYmznCglSmkBLSGx5lTyh0A/WtDcZuFLa9LakpSSZBu3EvPpZYWlCN86T4pW2gjqFIV3A7gpNbSVKZjrV8WhwaKfDA0SFL5ike5OgPSqDKu70e6SFFfiIWByq5ubm996G/vXLcrzKlxwyytaFhQWlaT1Tr2r5DLXXz1+LHiGe2df+Jao5V9vA0Yl8CCpp2V4S0rU48Ak/M4ry9dJH6muZF1uUuQW2JDTqGBypdUoISQfv50sYVoya7qD0PmkknS9uBJSR67PT2rF7YzPG3ExJ0J5IcTzpWghSFb/AJh0r7Htq08qE6JZWPk4vUq2qxxmsMZsPKlwJi/i3CJKNpKD1GvX335Vc8Ov7OQQXXG1APMr5HNH9CPY0hLfj2bX1gT49uecaUQgKUQgE+29bHvTl4S4rNxqzvqubyFz5jiXHUtnaGgBoIB8/PZreko46ZVByz2cuZ8OGX7mrKcMmJxnKgOstlH+zzdf5cpodHEn+L6hvYNd3DjNl5EJdnvFu/Y+UWopRc7apfME7+l5pX+YyvqUq/I9RVzKQdbAOu1LfjVaZcOPG4h2BlX7cxkF5SUdDNg95EZXqCkFafRSRrvVZaMmiuSzTo1ztUW5QnQ7ElsofYWP3kLSFJP6GigKFbB8R/aQvTj3UwcYitMb8g7IdUvX5tppkUss2dGLcY8ay54hNsu8dWPTnD0Sy6pfiRlKPoVhSNnzWPWmYlQUNigBY2NVWcygB+KC8ZEptPzNwkJHK84PpCz/AA70dHpVnJAGzXi8Ukbqq6tWQcWZjLi8iRvFmWLkpEpsS30uIjuuOdS46UeIsgj6UpHRKR3rlbaYTH21CZhoUhpZfUSoNpcB5Fq/l2AD6bpxO2aIp+KUJCW2HlSOQDZW4oEcxJ+5rit2MtRRBXpBMfxUOAjoppZJ8PR7pBI6H0rkNT9K03282j2K93tjBRKXj1iWZKWnIMRUhtHifCzo+yU76+E8k6WnfUbB10r1M/KIfFiNbTbW4uHx4rinFLHyIJTsOLWrp0OwEjt3piWy3x4MRiOkFSI4UGif3Uk9h7eVc+U2WJkWPyrPM5gxJRyL0ffp9+tdVt2lhpYcMdeHi66dl75J9p5/ovstZyXNbhjF6w27tOWBmUl5YSvlbcKHCCpXmpOgdD1pssjY3uq5w/xiHh2OM2GE846y0pa+ZzuVKOz0HRI69hVkSpAGugrdsmnhLxGvRW45lL1m9ecphuTGdjvJCm3UFCwfMEaIr0qn8XMk/C2C3O5tBTk9xAiW6Ok6U/Kd+RlA9SVEHp2AJ8qgbJG/2cHFucE8aStRPgsOMIJ/gbdWhH+lIoqw8Nsf/CuB2THCoLXb4TbLix++4E/Or81En86KA6cwsFryTHJ1lvLCpEGY14bqAdKHXYUk+SknRB8iBVIwnLbnjt3Z4f588kXXXLaLwscrN5aT2+bsmSBrmQe5+ZOwejPqHy7GrLlVndtF/trNwguaJacH0qHZSSOqVDyUNEUBKr6t9qV+Z5jllpu2SRrTaWJseFEQuO+pWhGcLRWoujupGh01130rZu38R8BRyWt1eeY+39EWU8lq6Rk+iXT8j4A8l8qvc1IWPifgdymOW6fLFhuro5Xrfe45hvnQ7f4mkr/6SRWU8ekZJtdPBXk5bkz0TJXWrzFjP2tKlxm1R0L50oSkq2Obm7q67Gh01W1yzPJbNcbhGkXGJM+DtrLwCo6Gw444N85AVzcqR8xCQRpJ2RTRbhW1SlyGYcVSnk6W4lpJ8Qe5HcVlcC3qk/ELhxy8E8odLQKgn03317VPlH9FX2p49KVY8gunxN7iuXOHemoVvTLanx2wlIWUq22oJJSfpChrro9aq+OZ9f7nIsbEi4xGPiFO/FrDSBzBPhcoG1cvZZHybPbp3pvRYEGKwWIkOPHZJJLbTQSkk9+g6da0/ZdsCWkC2RQljq0nwE/4Z9U9On5UUo/KEq5vxioj8RL+LdHYmNR49y+EXK34ZLchrxW0IcR6aClBSfJQ96m8ZzK7SMog2K6sNtOPzZaWHko+SSw3zAaPZLiSAFD0IPnVvyK6YtYIfxd+nWm2MNg6XLcQ2AO5A5v6CqSriIi/rS1w1xGVkq0qPJcXmzCt7RPdXjLG1/8AbSd+tY5LHgjXJPtjByW+WrHLNIvN7uEe3wIyeZ599WkpHp7k+QHUnoKX+JQrnnuWRM+yGE/b7PB5jjdpkJ5XNqGjNfT5OKSdISfpSST1Nddm4cyrhdo+R8Q7o3kd3jq8SJES2UW63q9WWj9Sx/xF7PpqmKgEJHN3qBeCEBO9edFbUUAUUUUAEbHnUffLHZr7DVDvdqhXKOru1KYS6n9FA0UUBSFcGcOir8TH3b9jS99rPdXmED38PmKP9NaJ4d5VHHLbeLuWtI7alMxZJ191t7/OiigM/gTPfPjJfyPPVqhA/wDzo/uukyk6vXEfOrgk9FNouCYyFD3DSEn/AM0UUBJWLhNw+s8pM1nGosucnr8XcFKlvE+vO6VHfvVzDSAEgbCU9kjoBRRQHpRRRQBRRRQH/9k=",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────────────────────────────────────
# TRADUCTIONS / TRANSLATIONS
# ─────────────────────────────────────────────────────────────────────────────
TRANSLATIONS = {
    "FR": {
        "lang_label":        "🌐 Langue / Language",
        "settings":          "⚙️ Paramètres",
        "data_source_label": "Source de données",
        "data_demo":         "🎲 Données démo",
        "data_live":         "🔗 Pipeline en direct",
        "ecmwf_run_title":   "### 📅 Date du run ECMWF",
        "date_label":        "Date",
        "hour_utc":          "Heure UTC",
        "swh_source":        "Source SWH",
        "vars_title":        "## 📊 Variables à visualiser",
        "period_title":      "## 🗓️ Période",
        "date_from":         "📅 Du",
        "date_to":           "📅 Au",
        "hour_from":         "🕐 Heure début",
        "hour_to":           "🕐 Heure fin",
        "display_title":     "## 🎨 Options d'affichage",
        "show_markers":      "Marqueurs sur courbes",
        "show_thresh":       "Seuils d'alerte",
        "chart_type":        "Type principal",
        "chart_types":       ["Séries temporelles", "Barres", "Aire empilée"],
        "footer_copy":       "© 2026 LAOUROU M. DIANE",
        "grp_waves":         "🌊 Vagues",
        "grp_wind":          "💨 Vent",
        "grp_pres":          "🔴 Pression / Temp",
        "grp_other":         "🌫️ Autres",
        "kpi_swh":           "SWH Max",
        "kpi_wind":          "Vent Max",
        "kpi_gust":          "Rafale Max",
        "kpi_mslp":          "MSLP Min",
        "kpi_sst":           "SST Moy",
        "kpi_rain":          "Précip Max",
        "alert_danger":      "🔴 DANGER",
        "alert_caution":     "🟡 PRUDENCE",
        "alert_normal":      "🟢 NORMAL",
        "warn_danger":       "Conditions dangereuses attendues.",
        "warn_caution":      "Conditions modérées. Prudence recommandée.",
        "warn_none":         "Warning : Aucune. Conditions clémentes.",
        "tab_meteo":         "🌤️ Météo",
        "tab_mer":           "🌊 Mer",
        "tab_wind":          "💨 Vent & Rose",
        "tab_data":          "📋 Données brutes",
        "tab_export":        "💾 Exports",
        "select_var_hint":   "👈 Sélectionnez au moins une variable dans la barre latérale.",
        "wind_speed_title":  "Vitesses de vent",
        "wind_dir_title":    "Direction du vent 10m (°)",
        "wave_height_title": "Hauteurs de houle",
        "swell_dir_title":   "Direction & Hauteur Swell",
        "swell_per_title":   "Périodes de swell",
        "current_title":     "Courants marins",
        "corr_title":        "Matrice de Corrélation",
        "scatter_title":     "Nuage de points personnalisé",
        "axis_x":            "Axe X",
        "axis_y":            "Axe Y",
        "rows_cols":         "lignes × colonnes",
        "export_data_title": "Exporter les données filtrées",
        "export_png_title":  "Exporter les graphiques (PNG)",
        "export_csv":        "⬇️ Export CSV (;)",
        "export_xlsx":       "⬇️ Export Excel (.xlsx)",
        "export_json":       "⬇️ Export JSON",
        "export_png":        "⬇️ Graphique (PNG)",
        "export_png_warn":   "Export PNG non disponible (kaleido manquant)",
        "bulletin_title":    "### 📋 Bulletin de synthèse",
        "bulletin_dl":       "⬇️ Télécharger le bulletin (.txt)",
        "bulletin_textarea": "Bulletin texte",
        "spinner_live":      "⏳ Lancement du pipeline ECMWF/Copernicus...",
        "err_pipeline":      "❌ Erreur pipeline",
        "info_fallback":     "💡 Passage aux données de démonstration.",
        "header_title":      "Prévision Marine — Sème",
        "header_demo_badge": "🎲 DONNÉES DÉMO",
        "header_source":     "Source : ECMWF Open Data + Copernicus Marine",
        "header_updated":    "Mise à jour",
        "bul_header":        "BULLETIN DE PRÉVISION MARINE — SÈME (6.22°N, 2.63°E)",
        "bul_generated":     "Généré le",
        "bul_period":        "Période",
        "bul_alert":         "NIVEAU D'ALERTE",
        "bul_stats":         "STATISTIQUES CLÉS",
        "bul_swh":           "SWH    max",
        "bul_wind":          "Vent   max",
        "bul_gust":          "Rafale max",
        "bul_mslp":          "MSLP   min",
        "bul_sst":           "SST    moy",
        "bul_source":        "Source : ECMWF Open Data + Copernicus Marine Service",
        "bul_author":        "Auteur : LAOUROU MAKONDJOU DIANE",
        "wind_rose_title":   "Rose des Vents 10m",
        "time_label":        "Temps",
        "forecast_title":    "Prévisions — Sème",
        # Word corrigé
        "word_upload_title": "## 📄 Bulletin Word corrigé",
        "word_upload_help":  "Uploadez le .docx corrigé. T°air, Visibilité et Pluie remplaceront les données pipeline.",
        "word_upload_label": "Charger le bulletin Word corrigé (.docx)",
        "word_loaded_ok":    "✅ Bulletin chargé — corrections appliquées",
        "word_loaded_rows":  "pas de temps corrigés",
        "word_load_error":   "❌ Erreur lecture Word",
        "word_corrected_badge": "✏️ CORRIGÉ",
        "word_cols_corrected":  "Colonnes corrigées",
        "word_reset":        "🔄 Réinitialiser (données pipeline)",
    },
    "EN": {
        "lang_label":        "🌐 Language / Langue",
        "settings":          "⚙️ Settings",
        "data_source_label": "Data source",
        "data_demo":         "🎲 Demo data",
        "data_live":         "🔗 Live pipeline",
        "ecmwf_run_title":   "### 📅 ECMWF Run Date",
        "date_label":        "Date",
        "hour_utc":          "UTC Hour",
        "swh_source":        "SWH source",
        "vars_title":        "## 📊 Variables to display",
        "period_title":      "## 🗓️ Time Period",
        "date_from":         "📅 From",
        "date_to":           "📅 To",
        "hour_from":         "🕐 Start hour",
        "hour_to":           "🕐 End hour",
        "display_title":     "## 🎨 Display options",
        "show_markers":      "Markers on curves",
        "show_thresh":       "Alert thresholds",
        "chart_type":        "Main chart type",
        "chart_types":       ["Time series", "Bars", "Stacked area"],
        "footer_copy":       "© 2026 LAOUROU M. DIANE",
        "grp_waves":         "🌊 Waves",
        "grp_wind":          "💨 Wind",
        "grp_pres":          "🔴 Pressure / Temp",
        "grp_other":         "🌫️ Other",
        "kpi_swh":           "SWH Max",
        "kpi_wind":          "Wind Max",
        "kpi_gust":          "Gust Max",
        "kpi_mslp":          "MSLP Min",
        "kpi_sst":           "SST Avg",
        "kpi_rain":          "Rain Max",
        "alert_danger":      "🔴 DANGER",
        "alert_caution":     "🟡 CAUTION",
        "alert_normal":      "🟢 NORMAL",
        "warn_danger":       "Dangerous conditions expected.",
        "warn_caution":      "Moderate conditions. Caution advised.",
        "warn_none":         "Warning: None. Calm conditions.",
        "tab_meteo":         "🌤️ Weather",
        "tab_mer":           "🌊 Sea",
        "tab_wind":          "💨 Wind & Rose",
        "tab_data":          "📋 Raw Data",
        "tab_export":        "💾 Exports",
        "select_var_hint":   "👈 Please select at least one variable in the sidebar.",
        "wind_speed_title":  "Wind speeds",
        "wind_dir_title":    "Wind direction 10m (°)",
        "wave_height_title": "Wave heights",
        "swell_dir_title":   "Swell Direction & Height",
        "swell_per_title":   "Swell periods",
        "current_title":     "Marine currents",
        "corr_title":        "Correlation Matrix",
        "scatter_title":     "Custom scatter plot",
        "axis_x":            "X Axis",
        "axis_y":            "Y Axis",
        "rows_cols":         "rows × columns",
        "export_data_title": "Export filtered data",
        "export_png_title":  "Export charts (PNG)",
        "export_csv":        "⬇️ Export CSV (;)",
        "export_xlsx":       "⬇️ Export Excel (.xlsx)",
        "export_json":       "⬇️ Export JSON",
        "export_png":        "⬇️ Chart (PNG)",
        "export_png_warn":   "PNG export unavailable (kaleido missing)",
        "bulletin_title":    "### 📋 Summary bulletin",
        "bulletin_dl":       "⬇️ Download bulletin (.txt)",
        "bulletin_textarea": "Bulletin text",
        "spinner_live":      "⏳ Running ECMWF/Copernicus pipeline...",
        "err_pipeline":      "❌ Pipeline error",
        "info_fallback":     "💡 Switching to demo data.",
        "header_title":      "Marine Forecast — Sème",
        "header_demo_badge": "🎲 DEMO DATA",
        "header_source":     "Source: ECMWF Open Data + Copernicus Marine",
        "header_updated":    "Updated",
        "bul_header":        "MARINE FORECAST BULLETIN — SÈME (6.22°N, 2.63°E)",
        "bul_generated":     "Generated on",
        "bul_period":        "Period",
        "bul_alert":         "ALERT LEVEL",
        "bul_stats":         "KEY STATISTICS",
        "bul_swh":           "SWH    max",
        "bul_wind":          "Wind   max",
        "bul_gust":          "Gust   max",
        "bul_mslp":          "MSLP   min",
        "bul_sst":           "SST    avg",
        "bul_source":        "Source: ECMWF Open Data + Copernicus Marine Service",
        "bul_author":        "Author: LAOUROU MAKONDJOU DIANE",
        "wind_rose_title":   "Wind Rose 10m",
        "time_label":        "Time",
        "forecast_title":    "Forecast — Sème",
        # Corrected Word
        "word_upload_title": "## 📄 Corrected Word bulletin",
        "word_upload_help":  "Upload the corrected .docx. T°air, Visibility and Rain will override pipeline data.",
        "word_upload_label": "Load corrected Word bulletin (.docx)",
        "word_loaded_ok":    "✅ Bulletin loaded — corrections applied",
        "word_loaded_rows":  "corrected time steps",
        "word_load_error":   "❌ Error reading Word file",
        "word_corrected_badge": "✏️ CORRECTED",
        "word_cols_corrected":  "Corrected columns",
        "word_reset":        "🔄 Reset (pipeline data)",
    },
}

# ─────────────────────────────────────────────────────────────────────────────
# MÉTADONNÉES VARIABLES (bilingues)
# ─────────────────────────────────────────────────────────────────────────────
ALERT_SWH_WARNING  = 1.6
ALERT_SWH_DANGER   = 2.0
ALERT_WIND_WARNING = 15
ALERT_WIND_DANGER  = 20

# ── Bénin Terminal — seuils rafales (km/h) par hauteur ───────────────────
BT_THRESHOLDS = {
    "10m": {"green": 28,  "yellow": 49,  "orange": 74 },
    "22m": {"green": 33,  "yellow": 57,  "orange": 86 },
    "60m": {"green": 40,  "yellow": 70,  "orange": 107},
    "70m": {"green": 40,  "yellow": 72,  "orange": 109},
}

POINTS = {
    "seme":    {"label": "🌊 Sème",              "lat": 6.22, "lon": 2.63},
    "terminal":{"label": "⚓ Port de Cotonou",   "lat": 6.35, "lon": 2.43},
}

VAR_META = {
    "swh_m":         {"FR": {"label": "Hauteur significative des vagues (SWH)", "short": "SWH",          "group": "grp_waves"},
                      "EN": {"label": "Significant Wave Height (SWH)",          "short": "SWH",          "group": "grp_waves"},
                      "unit":"m",   "color":"#15aabf","icon":"🌊",
                      "thresholds":[{"value":1.6,"color":"rgba(245,159,0,0.3)","dash":"dash",   "name_FR":"Prudence 1.6m","name_EN":"Caution 1.6m"},
                                    {"value":2.0,"color":"rgba(224,49,49,0.5)","dash":"dashdot","name_FR":"Danger 2.0m",  "name_EN":"Danger 2.0m"}]},
    "sw1_ht_m":      {"FR": {"label": "Hauteur Swell 1", "short": "Swell 1","group":"grp_waves"},
                      "EN": {"label": "Swell 1 Height",  "short": "Swell 1","group":"grp_waves"},
                      "unit":"m","color":"#339af0","icon":"🌊","thresholds":[]},
    "sw2_ht_m":      {"FR": {"label": "Hauteur Swell 2", "short": "Swell 2","group":"grp_waves"},
                      "EN": {"label": "Swell 2 Height",  "short": "Swell 2","group":"grp_waves"},
                      "unit":"m","color":"#74c0fc","icon":"🌊","thresholds":[]},
    "sw1_period_s":  {"FR": {"label": "Période Swell 1",  "short": "Période Sw1","group":"grp_waves"},
                      "EN": {"label": "Swell 1 Period",   "short": "Period Sw1", "group":"grp_waves"},
                      "unit":"s","color":"#a5d8ff","icon":"⏱️","thresholds":[]},
    "sw2_period_s":  {"FR": {"label": "Période Swell 2",  "short": "Période Sw2","group":"grp_waves"},
                      "EN": {"label": "Swell 2 Period",   "short": "Period Sw2", "group":"grp_waves"},
                      "unit":"s","color":"#d0ebff","icon":"⏱️","thresholds":[]},
    "wind10_spd_kt": {"FR": {"label": "Vitesse vent 10m","short":"Vent 10m", "group":"grp_wind"},
                      "EN": {"label": "Wind speed 10m",  "short":"Wind 10m", "group":"grp_wind"},
                      "unit":"kt","color":"#69db7c","icon":"💨",
                      "thresholds":[{"value":15,"color":"rgba(245,159,0,0.3)","dash":"dash",   "name_FR":"15 kt","name_EN":"15 kt"},
                                    {"value":20,"color":"rgba(224,49,49,0.5)","dash":"dashdot","name_FR":"20 kt","name_EN":"20 kt"}]},
    "wind10_gust_kt":{"FR": {"label": "Rafales 10m",    "short":"Rafales",  "group":"grp_wind"},
                      "EN": {"label": "Gusts 10m",       "short":"Gusts",    "group":"grp_wind"},
                      "unit":"kt","color":"#b2f2bb","icon":"💨","thresholds":[]},
    "wind100_spd_kt":{"FR": {"label": "Vitesse vent 100m","short":"Vent 100m","group":"grp_wind"},
                      "EN": {"label": "Wind speed 100m", "short":"Wind 100m","group":"grp_wind"},
                      "unit":"kt","color":"#40c057","icon":"💨","thresholds":[]},
    "wind10_dir":    {"FR": {"label": "Direction vent 10m","short":"Dir Vent 10m","group":"grp_wind"},
                      "EN": {"label": "Wind direction 10m","short":"Wind Dir 10m","group":"grp_wind"},
                      "unit":"°","color":"#a9e34b","icon":"🧭","thresholds":[]},
    "mslp_hpa":      {"FR": {"label": "Pression mer (MSLP)","short":"MSLP","group":"grp_pres"},
                      "EN": {"label": "Sea level pressure", "short":"MSLP","group":"grp_pres"},
                      "unit":"hPa","color":"#ffa94d","icon":"🔴",
                      "thresholds":[{"value":1010,"color":"rgba(21,170,191,0.2)","dash":"dot","name_FR":"1010 hPa","name_EN":"1010 hPa"}]},
    "t2m_c":         {"FR": {"label": "Température 2m",  "short":"T 2m","group":"grp_pres"},
                      "EN": {"label": "Temperature 2m",  "short":"T 2m","group":"grp_pres"},
                      "unit":"°C","color":"#ff6b6b","icon":"🌡️","thresholds":[]},
    "sst_c":         {"FR": {"label": "Température surface mer (SST)","short":"SST","group":"grp_pres"},
                      "EN": {"label": "Sea surface temperature (SST)","short":"SST","group":"grp_pres"},
                      "unit":"°C","color":"#f06595","icon":"🌡️","thresholds":[]},
    "vis_km":        {"FR": {"label": "Visibilité",  "short":"Visibilité","group":"grp_other"},
                      "EN": {"label": "Visibility",  "short":"Visibility","group":"grp_other"},
                      "unit":"km","color":"#e599f7","icon":"👁️","thresholds":[]},
    "rain_pct":      {"FR": {"label": "Probabilité de précipitation","short":"Précip. (%)","group":"grp_other"},
                      "EN": {"label": "Precipitation probability",   "short":"Rain (%)",   "group":"grp_other"},
                      "unit":"%","color":"#4dabf7","icon":"🌧️",
                      "thresholds":[{"value":50,"color":"rgba(21,170,191,0.2)","dash":"dot","name_FR":"50%","name_EN":"50%"}]},
    "cur_spd_ms":    {"FR": {"label": "Vitesse courant marin (m/s)","short":"Courant", "group":"grp_other"},
                      "EN": {"label": "Current speed",          "short":"Current", "group":"grp_other"},
                      "unit":"m/s","color":"#cc5de8","icon":"🔄","thresholds":[]},
    "cur_dir":       {"FR": {"label": "Direction courant marin","short":"Dir Courant","group":"grp_other"},
                      "EN": {"label": "Current direction",      "short":"Cur Dir",   "group":"grp_other"},
                      "unit":"°","color":"#da77f2","icon":"🧭","thresholds":[]},
}

# ─────────────────────────────────────────────────────────────────────────────
# CSS
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
:root{--ocean-dark:#0a1628;--ocean-mid:#0d2240;--ocean-blue:#0e3a6e;--ocean-teal:#0b7285;
      --ocean-cyan:#15aabf;--alert-red:#e03131;--alert-yellow:#f59f00;--alert-green:#2f9e44;
      --text-light:#e9ecef;--text-muted:#adb5bd;--card-bg:rgba(13,34,64,0.85);--border:rgba(21,170,191,0.3);}
.stApp{background:linear-gradient(160deg,#0a1628 0%,#0d2240 50%,#051020 100%);}
.marine-header{background:linear-gradient(135deg,rgba(14,58,110,0.9),rgba(11,114,133,0.7));
  border:1px solid var(--border);border-radius:16px;padding:1.5rem 2rem;margin-bottom:1.5rem;
  display:flex;align-items:center;gap:1.5rem;box-shadow:0 8px 32px rgba(0,0,0,0.4);}
.marine-header h1{color:var(--text-light);font-size:1.6rem;font-weight:700;margin:0;}
.marine-header .subtitle{color:var(--ocean-cyan);font-size:0.85rem;letter-spacing:1.5px;text-transform:uppercase;font-weight:600;}
.kpi-card{background:var(--card-bg);border:1px solid var(--border);border-radius:12px;
  padding:1.1rem 1.2rem;text-align:center;box-shadow:0 4px 20px rgba(0,0,0,0.3);
  transition:transform 0.2s,box-shadow 0.2s;height:110px;display:flex;flex-direction:column;justify-content:center;}
.kpi-card:hover{transform:translateY(-2px);box-shadow:0 8px 32px rgba(21,170,191,0.2);}
.kpi-card .kpi-label{font-size:0.72rem;color:var(--ocean-cyan);text-transform:uppercase;letter-spacing:1px;font-weight:600;margin-bottom:0.3rem;}
.kpi-card .kpi-value{font-size:1.4rem;font-weight:800;color:var(--text-light);line-height:1.2;}
.kpi-card .kpi-unit{font-size:0.8rem;color:var(--text-muted);margin-top:0.1rem;}
.warning-box{border-radius:10px;padding:1rem 1.4rem;margin:1rem 0;border-left:5px solid;font-size:0.9rem;font-weight:500;}
.warning-none{background:rgba(47,158,68,0.12);border-color:var(--alert-green);color:#a9e34b;}
.warning-yellow{background:rgba(245,159,0,0.12);border-color:var(--alert-yellow);color:#ffd43b;}
.warning-red{background:rgba(224,49,49,0.14);border-color:var(--alert-red);color:#ff8787;}
.section-title{color:var(--ocean-cyan);font-size:0.75rem;text-transform:uppercase;letter-spacing:2px;
  font-weight:700;border-bottom:1px solid var(--border);padding-bottom:0.5rem;margin:1.5rem 0 1rem 0;}
[data-testid="stSidebar"]{background:rgba(10,22,40,0.97)!important;border-right:1px solid var(--border);}
[data-testid="stSidebar"] .stMarkdown h2{color:var(--ocean-cyan);font-size:0.85rem;text-transform:uppercase;letter-spacing:1.5px;}
[data-testid="stMetric"]{background:var(--card-bg)!important;border:1px solid var(--border)!important;border-radius:10px!important;padding:0.8rem!important;}
[data-testid="stMetricLabel"]{color:var(--ocean-cyan)!important;font-size:0.75rem!important;}
[data-testid="stMetricValue"]{color:var(--text-light)!important;}
[data-testid="stTabs"] button{color:var(--text-muted)!important;}
[data-testid="stTabs"] button[aria-selected="true"]{color:var(--ocean-cyan)!important;border-bottom-color:var(--ocean-cyan)!important;}
.stDownloadButton>button{background:linear-gradient(135deg,var(--ocean-teal),var(--ocean-blue))!important;
  color:white!important;border:1px solid var(--ocean-cyan)!important;border-radius:8px!important;font-weight:600!important;}
.stDownloadButton>button:hover{background:linear-gradient(135deg,var(--ocean-cyan),var(--ocean-teal))!important;
  box-shadow:0 4px 16px rgba(21,170,191,0.4)!important;}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS TRADUCTION
# ─────────────────────────────────────────────────────────────────────────────
def T(key: str) -> str:
    lang = st.session_state.get("lang", "FR")
    return TRANSLATIONS[lang].get(key, key)

def VM(var_key: str, field: str) -> str:
    lang = st.session_state.get("lang", "FR")
    return VAR_META[var_key][lang].get(field, "")

def thresh_name(th: dict) -> str:
    lang = st.session_state.get("lang", "FR")
    return th.get(f"name_{lang}", th.get("name_FR", ""))


# ─────────────────────────────────────────────────────────────────────────────
# DONNÉES DÉMO
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_data(ttl=3600)
def generate_demo_data() -> pd.DataFrame:
    np.random.seed(42)
    now   = datetime.now().replace(minute=0, second=0, microsecond=0)
    times = [now - timedelta(hours=6*i) for i in range(9, -21, -1)]
    n     = len(times)
    t     = np.linspace(0, 2*np.pi, n)
    swh   = np.clip(1.2 + 0.6*np.sin(t) + 0.3*np.sin(2.3*t) + np.random.normal(0,0.08,n), 0.3, 3.5)
    df = pd.DataFrame({
        "valid_local":    times,
        "swh_m":          swh,
        "sw1_ht_m":       np.clip(swh*0.65 + np.random.normal(0,0.05,n), 0, 2.5),
        "sw1_period_s":   11 + 3*np.sin(t*0.7) + np.random.normal(0,0.3,n),
        "sw1_dir":        200 + 15*np.sin(t*0.4) + np.random.normal(0,3,n),
        "sw2_ht_m":       np.clip(swh*0.35 + np.random.normal(0,0.04,n), 0, 1.5),
        "sw2_period_s":   7 + 2*np.sin(t*1.1) + np.random.normal(0,0.3,n),
        "sw2_dir":        240 + 10*np.sin(t*0.6) + np.random.normal(0,3,n),
        "wind10_spd_kt":  np.clip(8 + 6*np.sin(t*1.2) + np.random.normal(0,0.5,n), 0, 30),
        "wind10_gust_kt": np.clip(12 + 8*np.sin(t*1.2) + np.random.normal(0,0.7,n), 0, 38),
        "wind10_dir":     190 + 20*np.sin(t*0.5) + np.random.normal(0,5,n),
        "wind100_spd_kt": np.clip(12 + 7*np.sin(t*1.1) + np.random.normal(0,0.6,n), 0, 35),
        "wind100_dir":    195 + 18*np.sin(t*0.5) + np.random.normal(0,4,n),
        "mslp_hpa":       1013 - 2*np.sin(t*0.8) + np.random.normal(0,0.3,n),
        "t2m_c":          29 + 2*np.sin(t*0.3) + np.random.normal(0,0.2,n),
        "sst_c":          28 + 1.5*np.sin(t*0.25) + np.random.normal(0,0.15,n),
        "vis_km":         np.clip(15 - 3*np.sin(t*0.9) + np.random.normal(0,0.5,n), 3, 20),
        "rain_pct":       np.clip(20 + 30*np.sin(t*0.8)**2 + np.random.normal(0,3,n), 0, 100),
        "cur_spd_ms":     np.clip(0.5 + 0.4*np.sin(t*1.3) + np.random.normal(0,0.05,n), 0, 2),
        "cur_dir":        150 + 30*np.sin(t*0.7) + np.random.normal(0,5,n),
    })
    df["valid_local"] = pd.to_datetime(df["valid_local"])
    return clean_df(df)


# ─────────────────────────────────────────────────────────────────────────────
# PIPELINE LIVE  (ne nécessite pas xarray côté dashboard)
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_data(ttl=3600, show_spinner=False)
def load_pipeline_data(run_date, run_hour, swh_source):
    try:
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        import config, extractor
        from datetime import datetime as _dt
        run_dt = _dt.strptime(f"{run_date} {run_hour:02d}:00", "%Y-%m-%d %H:%M")
        if swh_source:
            config.SWH_SOURCE = swh_source
        df_ecmwf = extractor.extract_ecmwf(run_dt)
        df_cop   = extractor.extract_copernicus(run_dt)
        df       = extractor.merge_sources(df_ecmwf, df_cop)
        df["valid_local"] = pd.to_datetime(df["valid_local"])
        return clean_df(df), None
    except Exception as e:
        return None, str(e)


@st.cache_data(ttl=1800, show_spinner=False)
def load_github_csv():
    """Cherche le dernier bulletin CSV sur GitHub via git tree API."""
    try:
        import urllib.request, json, io as _io
        with urllib.request.urlopen(GITHUB_TREE_URL, timeout=15) as resp:
            tree = json.loads(resp.read().decode("utf-8"))
        csv_files = sorted([
            item["path"] for item in tree.get("tree", [])
            if item["path"].startswith("bulletin_marine_seme_") and item["path"].endswith(".csv")
        ])
        if not csv_files:
            # Fallback ancien nom
            has_latest = any(item["path"] == "latest_forecast.csv" for item in tree.get("tree", []))
            if has_latest:
                latest = "latest_forecast.csv"
            else:
                return None, "Aucun CSV trouvé sur GitHub."
        else:
            latest = csv_files[-1]
        with urllib.request.urlopen(GITHUB_RAW_URL + latest, timeout=15) as resp:
            content = resp.read().decode("utf-8")
        df = pd.read_csv(_io.StringIO(content))
        df["valid_local"] = pd.to_datetime(df["valid_local"])
        return clean_df(df), f"📄 {latest}"
    except Exception as e:
        return None, str(e)


# ─────────────────────────────────────────────────────────────────────────────
# LECTURE DU WORD CORRIGÉ
# ─────────────────────────────────────────────────────────────────────────────
# Structure du Tableau 1 (index 1 dans doc.tables) :
#   row 0 : en-têtes groupes
#   row 1 : sous-en-têtes variables
#   row 2 : sous-en-têtes unités  → col 0=Date, 1=Time, 8=Vis.(km), 9=T(°C), 11=Rain(%)
#   rows 3+ : données

# Correspondance mois abrégés anglais → numéro
_MONTH_MAP = {"Jan":1,"Feb":2,"Mar":3,"Apr":4,"May":5,"Jun":6,
              "Jul":7,"Aug":8,"Sep":9,"Oct":10,"Nov":11,"Dec":12}
# Correspondance jour abrégé → (ignoré, on parse juste la date)

def _parse_word_datetime(date_str: str, time_str: str) -> pd.Timestamp | None:
    """Convertit 'Sat. 16 May' + '19:00' en Timestamp."""
    try:
        # Nettoyer : "Sat. 16 May" → ["16", "May"]
        parts = date_str.replace(".", "").split()
        # parts peut être ["Sat", "16", "May"] ou ["16", "May"]
        day_num = None
        month   = None
        for p in parts:
            if p.isdigit():
                day_num = int(p)
            elif p in _MONTH_MAP:
                month = _MONTH_MAP[p]
        if day_num is None or month is None:
            return None
        year  = datetime.now().year
        hour  = int(time_str.split(":")[0])
        minute= int(time_str.split(":")[1])
        return pd.Timestamp(year=year, month=month, day=day_num, hour=hour, minute=minute)
    except Exception:
        return None


def read_word_corrections(docx_bytes: bytes) -> tuple[pd.DataFrame | None, str | None]:
    """
    Lit le bulletin Word corrigé et extrait T(°C), Vis.(km), Rain(%)
    depuis le Tableau 1.
    Retourne (df_corrections, error_msg).
    df_corrections colonnes : valid_local, t2m_c, vis_km, rain_pct
    """
    try:
        from docx import Document
        import io as _io
        doc   = Document(_io.BytesIO(docx_bytes))
        table = doc.tables[1]   # Tableau principal (28 lignes × 23 colonnes)

        records = []
        prev_date = ""
        for row in table.rows[3:]:   # Sauter les 3 lignes d'en-têtes
            cells    = [c.text.strip() for c in row.cells]
            if len(cells) < 12:
                continue
            date_str = cells[0] if cells[0] else prev_date
            if cells[0]:
                prev_date = cells[0]
            time_str = cells[1]
            t_str    = cells[9]
            vis_str  = cells[8]
            rain_str = cells[11]

            ts = _parse_word_datetime(date_str, time_str)
            if ts is None:
                continue

            def _to_float(s):
                try:    return float(s.replace(",", "."))
                except: return np.nan

            records.append({
                "valid_local": ts,
                "t2m_c":       _to_float(t_str),
                "vis_km":      _to_float(vis_str),
                "rain_pct":    _to_float(rain_str),
            })

        if not records:
            return None, "Aucune donnée trouvée dans le tableau Word."

        df = pd.DataFrame(records)
        df = df.dropna(subset=["valid_local"])
        return df, None

    except Exception as e:
        return None, str(e)


def apply_word_corrections(df_base: pd.DataFrame, df_corr: pd.DataFrame) -> pd.DataFrame:
    """
    Fusionne les corrections Word dans le DataFrame de base.
    Stratégie : merge sur valid_local (tolérance ±30 min),
    les 3 colonnes corrigées remplacent les valeurs pipeline.
    """
    df = df_base.copy()
    df["valid_local"] = pd.to_datetime(df["valid_local"])
    df_corr = df_corr.copy()
    df_corr["valid_local"] = pd.to_datetime(df_corr["valid_local"])

    for _, corr_row in df_corr.iterrows():
        ts = corr_row["valid_local"]
        # Trouver la ligne la plus proche (tolérance 30 min)
        delta = (df["valid_local"] - ts).abs()
        idx   = delta.idxmin()
        if delta[idx] <= pd.Timedelta("30min"):
            for col in ["t2m_c", "vis_km", "rain_pct"]:
                val = corr_row[col]
                if not np.isnan(val):
                    df.at[idx, col] = val

    return df


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS GRAPHIQUES
# ─────────────────────────────────────────────────────────────────────────────
def get_alert_level(df):
    swh_max  = df["swh_m"].max()          if "swh_m"          in df.columns else 0
    # Alerte vent basée sur les RAFALES 10m (pas le vent moyen)
    gust_max = df["wind10_gust_kt"].max() if "wind10_gust_kt" in df.columns else 0
    if swh_max >= ALERT_SWH_DANGER or gust_max >= ALERT_WIND_DANGER:
        return T("alert_danger"),  "warning-red",    f"{T('warn_danger')}  SWH max {swh_max:.1f} m — Rafales max {gust_max:.0f} kt."
    elif swh_max >= ALERT_SWH_WARNING or gust_max >= ALERT_WIND_WARNING:
        return T("alert_caution"), "warning-yellow", f"{T('warn_caution')} SWH max {swh_max:.1f} m — Rafales max {gust_max:.0f} kt."
    else:
        return T("alert_normal"),  "warning-none",   f"{T('warn_none')}    SWH max {swh_max:.1f} m — Rafales max {gust_max:.0f} kt."


def clean_df(df: pd.DataFrame) -> pd.DataFrame:
    """
    Force toutes les colonnes numériques en float64 natif numpy.
    Nécessaire pour pandas 3.x / ArrowDtype sur Streamlit Cloud.
    """
    df = df.copy()
    for col in df.columns:
        if col == "valid_local":
            df[col] = pd.to_datetime(df[col])
        else:
            df[col] = pd.to_numeric(df[col], errors="coerce").astype("float64")
    return df


def plotly_theme():
    return dict(
        paper_bgcolor="rgba(10,22,40,0)", plot_bgcolor="rgba(13,34,64,0.5)",
        font=dict(color="#e9ecef", family="monospace,sans-serif", size=12),
        xaxis=dict(
            gridcolor="rgba(21,170,191,0.12)",
            linecolor="rgba(21,170,191,0.3)",
            showgrid=True,
            dtick=3 * 3600 * 1000,      # pas de 3h en millisecondes
            tickformat="%H:%M\n%d/%m",   # heure locale telle quelle
            tickangle=0,
            tickfont=dict(size=9),
            type="date",
        ),
        yaxis=dict(gridcolor="rgba(21,170,191,0.12)", linecolor="rgba(21,170,191,0.3)", showgrid=True),
        legend=dict(bgcolor="rgba(10,22,40,0.7)", bordercolor="rgba(21,170,191,0.3)", borderwidth=1),
        margin=dict(l=60,r=30,t=40,b=60), hovermode="x unified", transition={"duration": 0},
    )

def fig_to_bytes(fig):
    try:    return fig.to_image(format="png", scale=2, width=1400, height=600)
    except: return fig.to_html().encode()

def df_to_excel_bytes(df):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        df.to_excel(w, index=False)
    return buf.getvalue()

def df_to_csv_bytes(df):
    return df.to_csv(index=False, sep=";", decimal=",").encode("utf-8-sig")


def make_timeseries(df, selected_vars, title=""):
    lang = st.session_state.get("lang", "FR")
    n    = len(selected_vars)
    if n == 0:
        return go.Figure()

    # Convertir valid_local en string pour que Plotly affiche
    # l'heure locale exacte sans conversion UTC
    df = df.copy()
    x_vals = df["valid_local"].dt.strftime("%Y-%m-%d %H:%M:%S")

    fig = make_subplots(rows=n, cols=1, shared_xaxes=True, vertical_spacing=0.03,
        subplot_titles=[VAR_META.get(v,{}).get(lang,{}).get("label",v) for v in selected_vars])
    for i, var in enumerate(selected_vars, 1):
        if var not in df.columns: continue
        meta  = VAR_META.get(var, {})
        color = meta.get("color","#15aabf")
        unit  = meta.get("unit","")
        short = meta.get(lang,{}).get("short", var)
        r,g,b = int(color[1:3],16),int(color[3:5],16),int(color[5:7],16)
        chart_t = st.session_state.get("chart_type_val", "lines")
        if chart_t == "bars":
            fig.add_trace(go.Bar(
                x=x_vals, y=df[var], name=short,
                marker_color=f"rgba({r},{g},{b},0.8)",
                hovertemplate=f"<b>{short}</b>: %{{y:.2f}} {unit}<extra></extra>",
            ), row=i, col=1)
        else:
            fig.add_trace(go.Scattergl(
                x=x_vals, y=df[var],
                mode="lines+markers" if chart_t != "area" else "lines",
                name=short,
                line=dict(color=color,width=2), marker=dict(size=4,color=color),
                fill="tozeroy" if (i==1 or chart_t=="area") else "none",
                fillcolor=f"rgba({r},{g},{b},0.12)",
                hovertemplate=f"<b>{short}</b>: %{{y:.2f}} {unit}<extra></extra>",
            ), row=i, col=1)
        for th in meta.get("thresholds",[]):
            fig.add_hline(y=th["value"], line_dash=th["dash"], line_color=th["color"], line_width=1.5,
                annotation_text=thresh_name(th), annotation_font=dict(color=th["color"],size=10),
                annotation_bgcolor="rgba(10,22,40,0.6)", row=i, col=1)
        fig.update_yaxes(title_text=unit, row=i, col=1, title_font=dict(size=10,color=color))

        # Configuration spéciale pour MSLP : axe Y fixe 1000-1020 hPa, pas de 2
        if var == "mslp_hpa" and not df[var].dropna().empty:
            data_min = df[var].dropna().min()
            data_max = df[var].dropna().max()
            y_min = min(1000, int(data_min) - 2)
            y_max = max(1020, int(data_max) + 2)
            fig.update_yaxes(
                range=[y_min, y_max],
                dtick=2,
                tick0=y_min,
                row=i, col=1
            )

        # Configuration spéciale pour T°air et SST : pas de 2°C, adapté aux données
        if var in ["t2m_c", "sst_c"] and not df[var].dropna().empty:
            data_min = df[var].dropna().min()
            data_max = df[var].dropna().max()
            y_min = int(data_min) - 1
            y_max = int(data_max) + 2
            fig.update_yaxes(
                range=[y_min, y_max],
                dtick=2,
                tick0=y_min,
                row=i, col=1
            )
    th = plotly_theme()
    for i in range(1, n+1):
        ax = f"xaxis{'' if i==1 else i}"
        fig.update_layout(**{ax: th["xaxis"]})
        ay = f"yaxis{'' if i==1 else i}"
        fig.update_layout(**{ay: dict(gridcolor="rgba(21,170,191,0.12)",linecolor="rgba(21,170,191,0.3)",showgrid=True)})
    fig.update_layout(paper_bgcolor=th["paper_bgcolor"], plot_bgcolor=th["plot_bgcolor"],
        font=th["font"], legend=th["legend"], margin=th["margin"], hovermode="x unified", transition={"duration": 0},
        height=max(180*n,280),
        title=dict(text=title,font=dict(color="#15aabf",size=14)) if title else None)
    return fig


def make_wind_rose(df, dir_col="wind10_dir", spd_col="wind10_spd_kt", title=None):
    if dir_col not in df.columns or spd_col not in df.columns:
        return go.Figure()
    lang_r = st.session_state.get("lang","FR")
    if lang_r == "FR":
        labels = ["N","NNE","NE","ENE","E","ESE","SE","SSE","S","SSO","SO","OSO","O","ONO","NO","NNO"]
    else:
        labels = ["N","NNE","NE","ENE","E","ESE","SE","SSE","S","SSW","SW","WSW","W","WNW","NW","NNW"]
    bins   = np.arange(0, 361, 22.5, dtype=float)
    dirs   = df[dir_col].dropna().astype(float).to_numpy()
    speeds = df[spd_col].dropna().astype(float).to_numpy()
    ml     = min(len(dirs), len(speeds))
    dirs, speeds = dirs[:ml], speeds[:ml]
    sbins   = [0, 5, 10, 15, 20, 100]
    slabels = ["0–5 kt","5–10 kt","10–15 kt","15–20 kt",">20 kt"]
    colors  = ["#74c0fc","#15aabf","#69db7c","#ffa94d","#ff6b6b"]
    fig = go.Figure()
    for j, (smin, smax) in enumerate(zip(sbins[:-1], sbins[1:])):
        mask   = (speeds >= float(smin)) & (speeds < float(smax))
        d      = dirs[mask]
        counts = [int(np.sum((d >= bins[k]) & (d < (bins[k+1] if k < 15 else 360.0))))
                  for k in range(16)]
        fig.add_trace(go.Barpolar(r=counts, theta=labels, name=slabels[j],
            marker_color=colors[j], marker_line_color="rgba(10,22,40,0.5)",
            marker_line_width=0.5, opacity=0.85))
    th = plotly_theme()
    fig.update_layout(
        polar=dict(bgcolor="rgba(13,34,64,0.6)",
            radialaxis=dict(showticklabels=True,ticks="",gridcolor="rgba(21,170,191,0.2)",tickfont=dict(color="#adb5bd",size=9)),
            angularaxis=dict(direction="clockwise",gridcolor="rgba(21,170,191,0.15)",tickfont=dict(color="#e9ecef",size=11))),
        paper_bgcolor=th["paper_bgcolor"],font=th["font"],legend=th["legend"],
        margin=dict(l=40,r=40,t=50,b=40),height=380,
        title=dict(text=title or T("wind_rose_title"),font=dict(color="#69db7c",size=13)))
    return fig


def make_swell_compass(df):
    lang = st.session_state.get("lang","FR")
    fig  = go.Figure()
    for sw,col,hkey in [("sw1_dir","#339af0","sw1_ht_m"),("sw2_dir","#74c0fc","sw2_ht_m")]:
        if sw not in df.columns or hkey not in df.columns: continue
        label = VAR_META[hkey][lang]["short"]
        fig.add_trace(go.Scatterpolar(r=df[hkey].fillna(0),theta=df[sw].fillna(0),
            mode="markers",name=label,
            marker=dict(color=col,size=8,opacity=0.8,line=dict(color="white",width=0.5)),
            hovertemplate="Dir: %{theta:.0f}°<br>Ht: %{r:.2f} m<extra></extra>"))
    th = plotly_theme()
    fig.update_layout(
        polar=dict(bgcolor="rgba(13,34,64,0.6)",
            radialaxis=dict(showticklabels=True,ticks="",gridcolor="rgba(21,170,191,0.2)",tickfont=dict(color="#adb5bd",size=9)),
            angularaxis=dict(direction="clockwise",rotation=90,gridcolor="rgba(21,170,191,0.15)",tickfont=dict(color="#e9ecef",size=10))),
        paper_bgcolor=th["paper_bgcolor"],font=th["font"],legend=th["legend"],
        margin=dict(l=40,r=40,t=50,b=40),height=380,
        title=dict(text=T("swell_dir_title"),font=dict(color="#339af0",size=13)))
    return fig


def make_correlation_heatmap(df, num_vars):
    lang  = st.session_state.get("lang","FR")
    avail = [v for v in num_vars if v in df.columns]
    if len(avail) < 2: return go.Figure()
    corr   = df[avail].corr().round(2)
    labels = [VAR_META.get(v,{}).get(lang,{}).get("short",v) for v in avail]
    fig = go.Figure(go.Heatmap(z=corr.values,x=labels,y=labels,
        colorscale=[[0,"#e03131"],[0.5,"rgba(13,34,64,0.5)"],[1,"#15aabf"]],
        zmin=-1,zmax=1,text=corr.values.round(2),texttemplate="%{text}",
        textfont=dict(size=10),colorbar=dict(tickfont=dict(color="#e9ecef"))))
    th = plotly_theme()
    fig.update_layout(paper_bgcolor=th["paper_bgcolor"],plot_bgcolor=th["plot_bgcolor"],
        font=th["font"],margin=dict(l=80,r=30,t=40,b=80),height=420,
        title=dict(text=T("corr_title"),font=dict(color="#15aabf",size=13)))
    return fig


# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
def render_sidebar():
    with st.sidebar:
        # Logo
        st.markdown(f"""
        <div style='text-align:center;padding:0.8rem 0 1rem 0;'>
            <div style='background:rgba(13,34,64,0.8);border:2px solid rgba(21,170,191,0.4);
                        border-radius:12px;padding:0.8rem;display:inline-block;'>
                <img src="{_LOGO_URI}" style='width:90px;height:90px;object-fit:contain;display:block;'/>
            </div>
            <div style='color:#15aabf;font-size:0.7rem;letter-spacing:2px;text-transform:uppercase;font-weight:700;margin-top:0.6rem;'>METEO-BENIN</div>
            <div style='color:#adb5bd;font-size:0.65rem;margin-top:0.2rem;'>DPROM / SPAM</div>
        </div>
        </div>""", unsafe_allow_html=True)

        # ── Sélecteur de langue (en premier) ──────────────────
        lang_choice = st.radio(
            "🌐 Langue / Language",
            ["🇫🇷 Français", "🇬🇧 English"],
            index=0 if st.session_state.get("lang","FR") == "FR" else 1,
            horizontal=True,
            key="lang_radio",
        )
        st.session_state["lang"] = "FR" if lang_choice.startswith("🇫🇷") else "EN"

        st.divider()

        # ── Info utilisateur + déconnexion ───────────────────
        user_name   = st.session_state.get("user_name","")
        user_role   = st.session_state.get("user_role","client")
        user_points = st.session_state.get("user_points",["seme"])
        st.markdown(f"""
        <div style='background:#1E2130;border-radius:8px;padding:8px 12px;margin-bottom:0.5rem;'>
            👤 <b style='font-size:0.8rem;'>{user_name}</b><br>
            <span style='color:#adb5bd;font-size:0.7rem;'>
                {"🔑 Admin" if user_role=="admin" else "👁️ Visualisation"}
            </span>
        </div>""", unsafe_allow_html=True)
        if st.button("🚪 Déconnexion", use_container_width=True, key="logout_btn"):
            for k in ["authenticated","user_name","user_role","user_points",
                      "username","point","df_session","bt_df","df_loaded"]:
                st.session_state.pop(k, None)
            st.rerun()

        st.divider()

        # ── Sélecteur de point (admin uniquement) ─────────────
        if len(user_points) > 1:
            point_choice = st.radio(
                "📍 Point de prévision",
                ["🌊 Sème", "⚓ Port de Cotonou"],
                horizontal=True,
                key="point_radio",
            )
            st.session_state["point"] = "terminal" if "Cotonou" in point_choice else "seme"
        else:
            # Client : afficher son point sans option de changement
            pt_label = "⚓ Port de Cotonou" if user_points[0]=="terminal" else "🌊 Sème"
            st.markdown(f"**📍 Point :** {pt_label}")

        st.divider()

        # ── Contenu sidebar selon le point actif ──────────────
        _active_point = st.session_state.get("point", "seme")

        if _active_point == "terminal":
            # Sidebar Bénin Terminal — géré dans render_benin_terminal()
            # Ici on n'affiche rien de spécifique à Sème
            run_date, run_hour, swh_source = None, 0, "ecmwf"
            selected_vars = []
            time_start = datetime.now()
            time_end   = datetime.now()
            show_markers = True
            show_thresholds = True
            chart_type = "lines"
            return {
                "run_date": run_date, "run_hour": run_hour,
                "swh_source": swh_source, "selected_vars": selected_vars,
                "time_start": time_start, "time_end": time_end,
                "show_markers": show_markers,
                "show_thresholds": show_thresholds,
                "chart_type": chart_type,
            }

        # ── Sidebar Sème/WAPCO uniquement ─────────────────────
        st.markdown(f"## {T('settings')}")

        # Source de données
        data_source = st.radio(T("data_source_label"),
            [T("data_demo"), T("data_live")], index=0)

        run_date, run_hour, swh_source = None, 0, "ecmwf"
        if data_source == T("data_live"):
            st.markdown(T("ecmwf_run_title"))
            run_date   = st.date_input(T("date_label"), value=datetime.utcnow().date())
            run_hour   = st.selectbox(T("hour_utc"), [0,6,12,18], index=2)
            swh_source = st.selectbox(T("swh_source"), ["ecmwf","copernicus"], index=0)

        st.divider()

        # Variables
        st.markdown(T("vars_title"))
        lang_v = st.session_state.get("lang","FR")
        selected_vars = []
        for gk in ["grp_waves","grp_wind","grp_pres","grp_other"]:
            vars_in_group = [k for k,m in VAR_META.items() if m[lang_v]["group"] == gk]
            with st.expander(T(gk), expanded=(gk=="grp_waves")):
                for v in vars_in_group:
                    meta    = VAR_META[v]
                    checked = v in ["swh_m","wind10_spd_kt","mslp_hpa"]
                    if st.checkbox(f"{meta['icon']} {meta[lang_v]['short']} ({meta['unit']})",
                                   value=checked, key=f"chk_{v}"):
                        selected_vars.append(v)

        st.divider()

        # Période — basée sur les vraies données GitHub si disponibles
        st.markdown(T("period_title"))

        # Bornes depuis session_state (df chargé dans main() avant sidebar)
        _df_b = st.session_state.get("df_loaded", generate_demo_data())

        # Bornes réelles
        _df_b["valid_local"] = pd.to_datetime(_df_b["valid_local"])
        _dt_min = _df_b["valid_local"].min().to_pydatetime()
        _dt_max = _df_b["valid_local"].max().to_pydatetime()
        _times  = sorted(_df_b["valid_local"].dt.to_pydatetime().tolist())

        # Toggle : afficher les échéances passées
        show_past = st.toggle(
            "🕐 Voir les échéances passées" if st.session_state.get("lang","fr")=="fr" else "🕐 Show past forecasts",
            value=False, key="show_past"
        )

        # Valeur par défaut : 19h du jour J (première échéance utile du bulletin)
        _19h_jour_j = _dt_min.replace(hour=19, minute=0, second=0, microsecond=0)
        # Si 19h est dans les données → partir de 19h, sinon premier pas disponible
        _19h_times = [t for t in _times if t >= _19h_jour_j]
        _def_start = _19h_times[0] if _19h_times and not show_past else _dt_min
        _def_end   = _dt_max

        col_a, col_b = st.columns(2)
        with col_a:
            start_date = st.date_input(T("date_from"), value=_def_start.date(),
                min_value=_dt_min.date(), max_value=_dt_max.date(), key="sd")
            _hs = sorted({t.hour for t in _times if t.date()==start_date}) or list(range(0,24,3))
            # Index par défaut = heure 19h si disponible
            _def_sh = _def_start.hour if start_date == _def_start.date() else _hs[0]
            _sh_idx = _hs.index(_def_sh) if _def_sh in _hs else 0
            start_hour = st.selectbox(T("hour_from"), _hs,
                format_func=lambda h: f"{h:02d}:00", index=_sh_idx, key="sh")
        with col_b:
            end_date = st.date_input(T("date_to"), value=_def_end.date(),
                min_value=_dt_min.date(), max_value=_dt_max.date(), key="ed")
            _he = sorted({t.hour for t in _times if t.date()==end_date}) or list(range(0,24,3))
            _def_eh = _def_end.hour if end_date == _def_end.date() else _he[-1]
            _eh_idx = _he.index(_def_eh) if _def_eh in _he else len(_he)-1
            end_hour = st.selectbox(T("hour_to"), _he,
                format_func=lambda h: f"{h:02d}:00", index=_eh_idx, key="eh")

        from datetime import datetime as _dt2
        time_start = _dt2.combine(start_date, _dt2.min.time()).replace(hour=start_hour)
        time_end   = _dt2.combine(end_date,   _dt2.min.time()).replace(hour=end_hour)
        if time_start > time_end:
            time_start, time_end = time_end, time_start

        st.divider()

        # ── Upload Word corrigé ────────────────────────────────
        st.markdown(T("word_upload_title"))
        st.caption(T("word_upload_help"))
        uploaded_word = st.file_uploader(
            T("word_upload_label"),
            type=["docx"],
            key="word_uploader",
        )
        if uploaded_word is not None:
            st.session_state["word_bytes"] = uploaded_word.read()
            st.session_state["word_name"]  = uploaded_word.name
        if st.session_state.get("word_bytes"):
            if st.button(T("word_reset"), width='stretch'):
                st.session_state.pop("word_bytes", None)
                st.session_state.pop("word_name",  None)
                st.rerun()

        st.divider()

        # Options d'affichage
        st.markdown(T("display_title"))
        show_markers    = st.checkbox(T("show_markers"), value=True)
        show_thresholds = st.checkbox(T("show_thresh"),  value=True)
        chart_type      = st.selectbox(T("chart_type"),  T("chart_types"))

        st.divider()
        st.markdown(f"""
        <div style='color:#adb5bd;font-size:0.65rem;text-align:center;line-height:1.6;'>
            Sème — 6.22°N, 2.63°E<br>Golfe de Guinée, Bénin<br>
            ECMWF · Copernicus<br>{T('footer_copy')}
        </div>""", unsafe_allow_html=True)

    # Stocker le type de graphique dans session_state AVANT le return
    if chart_type in ["Barres","Bars"]:
        st.session_state["chart_type_val"] = "bars"
    elif chart_type in ["Aire empilée","Stacked area"]:
        st.session_state["chart_type_val"] = "area"
    else:
        st.session_state["chart_type_val"] = "lines"

    return {"data_source":data_source,"run_date":str(run_date) if run_date else None,
            "run_hour":run_hour,"swh_source":swh_source,"selected_vars":selected_vars,
            "time_start":time_start,"time_end":time_end,
            "show_markers":show_markers,"show_thresholds":show_thresholds,"chart_type":chart_type}


# ─────────────────────────────────────────────────────────────────────────────
# KPI ROW
# ─────────────────────────────────────────────────────────────────────────────
def render_kpi_row(df):
    cols = st.columns(6)
    kpis = [
        ("swh_m",         T("kpi_swh"),  "m",   lambda s: f"{s.max():.2f}"),
        ("wind10_spd_kt", T("kpi_wind"), "kt",  lambda s: f"{s.max():.1f}"),
        ("wind10_gust_kt",T("kpi_gust"), "kt",  lambda s: f"{s.max():.1f}"),
        ("mslp_hpa",      T("kpi_mslp"), "hPa", lambda s: f"{s.min():.1f}"),
        ("sst_c",         T("kpi_sst"),  "°C",  lambda s: f"{s.mean():.1f}"),
        ("rain_pct",      T("kpi_rain"), "%",   lambda s: f"{s.max():.0f}"),
    ]
    for col,(var,label,unit,fmt) in zip(cols,kpis):
        with col:
            val = fmt(df[var].dropna()) if var in df.columns else "—"
            st.markdown(f"""
            <div class="kpi-card">
                <div class="kpi-label">{label}</div>
                <div class="kpi-value">{val}</div>
                <div class="kpi-unit">{unit}</div>
            </div>""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# ONGLETS
# ─────────────────────────────────────────────────────────────────────────────
@st.fragment
def render_main_tabs(df, df_filtered, params):
    lang = st.session_state.get("lang","FR")

    tab_meteo,tab_vent,tab_mer,tab_data,tab_export = st.tabs([
        T("tab_meteo"), T("tab_wind"), T("tab_mer"), T("tab_data"), T("tab_export")])

    # ── 🌤️ MÉTÉO ──────────────────────────────────────────────────────────────
    with tab_meteo:
        t_from = df_filtered["valid_local"].min().strftime("%d/%m")
        t_to   = df_filtered["valid_local"].max().strftime("%d/%m %H:%M")
        meteo_default = [v for v in ["mslp_hpa","t2m_c"] if v in df_filtered.columns]
        _opt_labels = {
            "sst_c":    "SST (°C)",
            "rain_pct": "Précip. (%)" if lang=="FR" else "Rain (%)",
            "vis_km":   "Visibilité (km)" if lang=="FR" else "Visibility (km)",
        }
        st.markdown("**Options :**" if lang=="FR" else "**Options:**")
        _cols = st.columns(3)
        meteo_extra = []
        for i,(var,label) in enumerate(_opt_labels.items()):
            if var in df_filtered.columns:
                if _cols[i].checkbox(label, value=False, key=f"meteo_{var}"):
                    meteo_extra.append(var)
        meteo_vars = meteo_default + meteo_extra
        if meteo_vars:
            st.plotly_chart(make_timeseries(df_filtered, meteo_vars,
                f"{T('forecast_title')} ({t_from} → {t_to})"), width='stretch')

    # ── 💨 VENT ───────────────────────────────────────────────────────────────
    with tab_vent:
        lang_a = st.session_state.get("lang","FR")
        def deg_to_card(deg):
            dirs = (["N","NNE","NE","ENE","E","ESE","SE","SSE","S","SSO","SO","OSO","O","ONO","NO","NNO"]
                    if lang_a=="FR" else
                    ["N","NNE","NE","ENE","E","ESE","SE","SSE","S","SSW","SW","WSW","W","WNW","NW","NNW"])
            try: return dirs[round(float(deg)/22.5) % 16]
            except: return ""

        wv = [v for v in ["wind10_spd_kt","wind10_gust_kt","wind100_spd_kt"] if v in df_filtered.columns]
        fig_wind = make_timeseries(df_filtered, wv, T("wind_speed_title"))
        _x_str = df_filtered["valid_local"].dt.strftime("%Y-%m-%d %H:%M:%S").tolist()
        if "wind10_dir" in df_filtered.columns and "wind10_spd_kt" in df_filtered.columns:
            _c10 = [deg_to_card(d) for d in df_filtered["wind10_dir"].tolist()]
            _y10 = df_filtered["wind10_spd_kt"].tolist()
            for idx,(_x,_y,_c) in enumerate(zip(_x_str,_y10,_c10)):
                if idx%2==0 and _c and pd.notna(_y):
                    fig_wind.add_annotation(x=_x, y=_y, text=f"<b>{_c}</b>",
                        showarrow=False, font=dict(size=8,color="#a9e34b"), yshift=12, row=1, col=1)
        if "wind100_dir" in df_filtered.columns and "wind100_spd_kt" in df_filtered.columns:
            _wv = [v for v in ["wind10_spd_kt","wind10_gust_kt","wind100_spd_kt"] if v in df_filtered.columns]
            _r100 = _wv.index("wind100_spd_kt")+1 if "wind100_spd_kt" in _wv else 3
            _c100 = [deg_to_card(d) for d in df_filtered["wind100_dir"].tolist()]
            _y100 = df_filtered["wind100_spd_kt"].tolist()
            for idx,(_x,_y,_c) in enumerate(zip(_x_str,_y100,_c100)):
                if idx%2==0 and _c and pd.notna(_y):
                    fig_wind.add_annotation(x=_x, y=_y, text=f"<b>{_c}</b>",
                        showarrow=False, font=dict(size=8,color="#40c057"), yshift=12, row=_r100, col=1)
        st.plotly_chart(fig_wind, width='stretch')

        lang_vv = st.session_state.get("lang","FR")
        title_100m = "Rose des Vents 100m" if lang_vv=="FR" else "Wind Rose 100m"
        c_r1,c_r2 = st.columns(2)
        with c_r1:
            st.plotly_chart(make_wind_rose(df_filtered, dir_col="wind10_dir",
                spd_col="wind10_spd_kt", title=T("wind_rose_title")),
                width='stretch', key="wind_rose_10m")
        with c_r2:
            if "wind100_dir" in df_filtered.columns and "wind100_spd_kt" in df_filtered.columns:
                st.plotly_chart(make_wind_rose(df_filtered, dir_col="wind100_dir",
                    spd_col="wind100_spd_kt", title=title_100m),
                    width='stretch', key="wind_rose_100m")

    # ── 🌊 MER ────────────────────────────────────────────────────────────────
    with tab_mer:
        lang_m = st.session_state.get("lang","FR")
        c1,c2 = st.columns(2)
        with c1:
            if "swh_m" in df_filtered.columns:
                st.plotly_chart(make_timeseries(df_filtered, ["swh_m"],
                    T("wave_height_title")), width='stretch')
        with c2:
            st.plotly_chart(make_swell_compass(df_filtered), width='stretch')

        _opt_mer = {
            "swell_ht":  "Hauteurs Swell 1 & 2" if lang_m=="FR" else "Swell 1 & 2 Heights",
            "swell_per": "Périodes Swell" if lang_m=="FR" else "Swell Periods",
            "currents":  "Courants marins" if lang_m=="FR" else "Marine Currents",
        }
        st.markdown("**Options :**" if lang_m=="FR" else "**Options:**")
        _cols_m = st.columns(3)
        show_sh  = _cols_m[0].checkbox(_opt_mer["swell_ht"],  value=False, key="mer_swell_ht")
        show_sp  = _cols_m[1].checkbox(_opt_mer["swell_per"], value=False, key="mer_swell_per")
        show_cur = _cols_m[2].checkbox(_opt_mer["currents"],  value=False, key="mer_currents")
        if show_sh:
            sv = [v for v in ["sw1_ht_m","sw2_ht_m"] if v in df_filtered.columns]
            if sv: st.plotly_chart(make_timeseries(df_filtered, sv, T("wave_height_title")), width='stretch')
        if show_sp:
            pv = [v for v in ["sw1_period_s","sw2_period_s"] if v in df_filtered.columns]
            if pv: st.plotly_chart(make_timeseries(df_filtered, pv, T("swell_per_title")), width='stretch')
        if show_cur:
            cv = [v for v in ["cur_spd_ms","cur_dir"] if v in df_filtered.columns]
            if cv: st.plotly_chart(make_timeseries(df_filtered, cv, T("current_title")), width='stretch')


    # Données brutes
    with tab_data:
        st.markdown(f"**{len(df_filtered)} {T('rows_cols')} {len(df_filtered.columns)}**")
        display_df = df_filtered.copy()
        display_df["valid_local"] = display_df["valid_local"].dt.strftime("%d/%m/%Y %H:%M")
        def color_swh(val):
            try:
                v = float(val)
                if v >= ALERT_SWH_DANGER:  return "background-color:rgba(224,49,49,0.25);color:#ff8787"
                if v >= ALERT_SWH_WARNING: return "background-color:rgba(245,159,0,0.2);color:#ffd43b"
            except Exception: pass
            return ""
        styled = display_df.style
        if "swh_m" in display_df.columns:
            styled = styled.map(color_swh, subset=["swh_m"])
        nc = display_df.select_dtypes(include=[np.number]).columns
        styled = styled.format({c:"{:.2f}" for c in nc}, na_rep="—")
        st.dataframe(styled, width='stretch', height=450)


    # Exports
    with tab_export:
        ts = datetime.now().strftime('%Y%m%d_%H%M')
        st.markdown(f'<div class="section-title">{T("export_data_title")}</div>', unsafe_allow_html=True)
        c1,c2,c3 = st.columns(3)
        with c1:
            st.download_button(T("export_csv"), data=df_to_csv_bytes(df_filtered),
                file_name=f"seme_{ts}.csv", mime="text/csv", width='stretch')
        with c2:
            st.download_button(T("export_xlsx"), data=df_to_excel_bytes(df_filtered),
                file_name=f"seme_{ts}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                width='stretch')
        with c3:
            json_str = df_filtered.assign(
                valid_local=df_filtered["valid_local"].dt.strftime("%Y-%m-%dT%H:%M:%S")
            ).to_json(orient="records", indent=2, force_ascii=False)
            st.download_button(T("export_json"), data=json_str.encode("utf-8"),
                file_name=f"seme_{ts}.json", mime="application/json", width='stretch')

        st.markdown(f'<div class="section-title">{T("export_png_title")}</div>', unsafe_allow_html=True)
        if params["selected_vars"]:
            try:
                img = fig_to_bytes(make_timeseries(df_filtered, params["selected_vars"]))
                ext = "png" if isinstance(img,bytes) and img[:4]==b'\x89PNG' else "html"
                st.download_button(T("export_png"), data=img,
                    file_name=f"seme_chart_{ts}.{ext}", mime=f"image/{ext}", width='stretch')
            except Exception as e:
                st.warning(f"{T('export_png_warn')} : {e}")

        st.divider()
        st.markdown(T("bulletin_title"))
        level,css,warning_txt = get_alert_level(df_filtered)
        swh_m  = df_filtered["swh_m"].max()          if "swh_m"          in df_filtered.columns else 0
        wind_m = df_filtered["wind10_spd_kt"].max()  if "wind10_spd_kt"  in df_filtered.columns else 0
        gust_m = df_filtered["wind10_gust_kt"].max() if "wind10_gust_kt" in df_filtered.columns else 0
        mslp_m = df_filtered["mslp_hpa"].min()       if "mslp_hpa"       in df_filtered.columns else 0
        sst_m  = df_filtered["sst_c"].mean()          if "sst_c"          in df_filtered.columns else 0
        t_from = df_filtered["valid_local"].min().strftime("%d/%m/%Y %H:%M")
        t_to   = df_filtered["valid_local"].max().strftime("%d/%m/%Y %H:%M")
        bulletin = (
            f"{T('bul_header')}\nMETEO-BENIN / DPROM / SPAM\n"
            f"{T('bul_generated')} : {datetime.now().strftime('%d/%m/%Y %H:%M')} (UTC+1)\n"
            f"{T('bul_period')}    : {t_from} → {t_to}\n\n"
            f"{'═'*55}\n{T('bul_alert')} : {level}\n{warning_txt}\n{'═'*55}\n\n"
            f"{T('bul_stats')}\n{'─'*17}\n"
            f"{T('bul_swh')}  : {swh_m:.2f} m\n"
            f"{T('bul_wind')} : {wind_m:.1f} kt\n"
            f"{T('bul_gust')} : {gust_m:.1f} kt\n"
            f"{T('bul_mslp')} : {mslp_m:.1f} hPa\n"
            f"{T('bul_sst')}  : {sst_m:.1f} °C\n\n"
            f"{'═'*55}\n{T('bul_source')}\n{T('bul_author')}\n"
        )
        st.text_area(T("bulletin_textarea"), value=bulletin, height=340)
        st.download_button(T("bulletin_dl"), data=bulletin.encode("utf-8"),
            file_name=f"bulletin_seme_{ts}.txt", mime="text/plain", width='stretch')



# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
def main():
    # ── Vérification authentification ─────────────────────────────────────
    if not st.session_state.get("authenticated", False):
        render_login()
        return

    user_role   = st.session_state.get("user_role", "client")
    user_name   = st.session_state.get("user_name", "")
    user_points = st.session_state.get("user_points", ["seme"])

    # Forcer le point si client avec un seul point
    if len(user_points) == 1:
        st.session_state["point"] = user_points[0]

    if "lang" not in st.session_state:
        st.session_state["lang"] = "FR"

    # ── Chargement une seule fois par session (session_state) ──────────
    is_demo   = False
    is_github = False
    csv_name  = ""

    if "df_session" not in st.session_state or "df_source" not in st.session_state:
        # Premier chargement uniquement
        df_gh, info_gh = load_github_csv()
        if df_gh is not None and not df_gh.empty:
            st.session_state["df_session"] = df_gh
            st.session_state["df_source"]  = "github"
            st.session_state["csv_name"]   = info_gh or ""
        else:
            st.session_state["df_session"] = generate_demo_data()
            st.session_state["df_source"]  = "demo"
            st.session_state["csv_name"]   = ""

    df       = st.session_state["df_session"]
    is_github = st.session_state["df_source"] == "github"
    is_demo   = st.session_state["df_source"] == "demo"
    csv_name  = st.session_state.get("csv_name", "")

    st.session_state["df_loaded"] = df

    params = render_sidebar()

    # ── Routage selon le point sélectionné ───────────────────
    if st.session_state.get("point") == "terminal":
        if "terminal" in user_points:
            render_benin_terminal()
        else:
            st.error("❌ Accès non autorisé à ce point.")
        return

    if "seme" not in user_points:
        st.error("❌ Accès non autorisé à ce point.")
        return

    # (données déjà chargées ci-dessus)

    # Filtre temporel
    df_filtered = df[
        (df["valid_local"] >= pd.Timestamp(params["time_start"])) &
        (df["valid_local"] <= pd.Timestamp(params["time_end"]))
    ].copy()
    if df_filtered.empty:
        df_filtered = df.copy()

    # ── Application des corrections Word ─────────────────────
    is_corrected = False
    word_info    = ""
    if st.session_state.get("word_bytes"):
        df_corr, err_word = read_word_corrections(st.session_state["word_bytes"])
        if err_word:
            st.error(f"{T('word_load_error')} : {err_word}")
        else:
            n_corrected  = len(df_corr)
            df           = apply_word_corrections(df, df_corr)
            df_filtered  = apply_word_corrections(df_filtered, df_corr)
            is_corrected = True
            word_info    = (f"{T('word_loaded_ok')} — "
                            f"**{n_corrected}** {T('word_loaded_rows')} "
                            f"· {T('word_cols_corrected')} : T°air, Visibilité, Pluie"
                            f" · *{st.session_state.get('word_name','')}*")
            st.success(word_info)

    # Header
    now_loc      = now_local()
    demo_badge   = (f" · <span style='color:#ffa94d;font-size:0.7rem;'>{T('header_demo_badge')}</span>"
                    if is_demo else "")
    github_badge = (f" · <span style='color:#69db7c;font-size:0.7rem;'>✅ PIPELINE</span>"
                    if is_github else "")
    corr_badge   = (f" · <span style='color:#69db7c;font-size:0.7rem;'>{T('word_corrected_badge')}</span>"
                    if is_corrected else "")
    st.markdown(f"""
    <div class="marine-header">
        <div style="font-size:3rem;">🌊</div>
        <div>
            <div class="subtitle">METEO-BENIN · DPROM / SPAM</div>
            <h1>{T('header_title')}{demo_badge}{github_badge}{corr_badge}</h1>
            <div style="color:#adb5bd;font-size:0.78rem;margin-top:0.2rem;">
                📍 6.22°N, 2.63°E · Golfe de Guinée, Bénin &nbsp;|&nbsp;
                {T('header_source')} &nbsp;|&nbsp;
                {T('header_updated')} : {now_loc.strftime('%d/%m/%Y %H:%M')} (UTC+1)
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

    # Warning banner
    level, css, warning_txt = get_alert_level(df_filtered)
    st.markdown(f'<div class="warning-box {css}"><b>{level}</b> — {warning_txt}</div>',
                unsafe_allow_html=True)

    # KPI
    render_kpi_row(df_filtered)
    st.markdown("---")

    # Onglets
    render_main_tabs(df, df_filtered, params)

    # Footer
    st.markdown(f"""
    <div style='text-align:center;color:#4a6480;font-size:0.68rem;margin-top:2rem;padding:1rem 0;
                border-top:1px solid rgba(21,170,191,0.15);'>
        © 2026 · LAOUROU MAKONDJOU DIANE · Météorologiste &amp; Data Scientist · METEO-BENIN / DPROM / SPAM<br>
        ECMWF Open Data (CC BY 4.0) · Copernicus Marine Service
    </div>""", unsafe_allow_html=True)


# =============================================================================
# BÉNIN TERMINAL — Fonctions dédiées
# =============================================================================

def bt_get_alert(val, height):
    """Niveau d'alerte rafale pour une hauteur donnée (Bénin Terminal)."""
    t = BT_THRESHOLDS[height]
    if val <= t["green"]:   return "green",  "🟢 Vert"
    if val <= t["yellow"]:  return "yellow", "🟡 Jaune"
    if val <= t["orange"]:  return "orange", "🟠 Orange"
    return "red", "🔴 Rouge"

def bt_global_alert(df):
    """Niveau d'alerte global Bénin Terminal."""
    levels = {"green":0,"yellow":1,"orange":2,"red":3}
    mx = 0
    for h, col in [("10m","RafaleV10_Km/h"),("22m","RafaleV22_Km/h"),
                   ("60m","RafaleV60_Km/h"),("70m","RafaleV70_Km/h")]:
        if col in df.columns:
            for v in df[col].dropna():
                lvl, _ = bt_get_alert(v, h)
                mx = max(mx, levels[lvl])
    labels = {
        0: ("🟢 Vert",   "warning-none",   "Conditions favorables aux opérations portuaires."),
        1: ("🟡 Jaune",  "warning-yellow",  "Vigilance recommandée — surveiller l'évolution."),
        2: ("🟠 Orange", "warning-yellow",  "Conditions difficiles — opérations à évaluer."),
        3: ("🔴 Rouge",  "warning-red",     "Opérations dangereuses — arrêt recommandé."),
    }
    return labels[mx]

def bt_load_csv():
    """Charge le dernier CSV Bénin Terminal depuis GitHub."""
    import requests, io
    url = f"https://api.github.com/repos/{GITHUB_OWNER}/{GITHUB_REPO_BT}/git/trees/main?recursive=1"
    try:
        r = requests.get(url, timeout=10)
        files = [f["path"] for f in r.json().get("tree",[])
                 if f["path"].endswith(".csv") and "ECMWF_Port" in f["path"]]
        if not files: return None, []
        files = sorted(files, reverse=True)
        raw = f"https://raw.githubusercontent.com/{GITHUB_OWNER}/{GITHUB_REPO_BT}/main/{files[0]}"
        df = pd.read_csv(io.StringIO(requests.get(raw, timeout=10).text))
        df["forecast_time_local"] = pd.to_datetime(df["forecast_time_local"])
        return df, files
    except Exception as e:
        return None, []

def bt_plot_wind(df, height, v_col, g_col):
    """Graphique vent + rafales avec zones colorées pour Bénin Terminal."""
    import plotly.graph_objects as go
    t = BT_THRESHOLDS[height]
    x = df["forecast_time_local"].dt.strftime("%a %d/%m %Hh")
    fig = go.Figure()
    ymax = max(df[g_col].max() * 1.2 if not df[g_col].empty else 50, t["orange"] + 20)
    zones = [
        (0, t["green"],  "rgba(46,204,113,0.08)"),
        (t["green"],  t["yellow"], "rgba(241,196,15,0.10)"),
        (t["yellow"], t["orange"], "rgba(230,126,34,0.12)"),
        (t["orange"], ymax,        "rgba(231,76,60,0.14)"),
    ]
    for y0, y1, col in zones:
        fig.add_hrect(y0=y0, y1=y1, fillcolor=col, line_width=0)
    for val, color in [(t["green"],"#2ECC71"),(t["yellow"],"#F1C40F"),(t["orange"],"#E67E22")]:
        fig.add_hline(y=val, line_dash="dot", line_color=color, line_width=1, opacity=0.6)
    fig.add_trace(go.Scatter(x=x, y=df[v_col], name=f"Vent {height}",
        line=dict(color="#4FC3F7", width=2), mode="lines+markers", marker=dict(size=4)))
    fig.add_trace(go.Scatter(x=x, y=df[g_col], name=f"Rafales {height}",
        line=dict(color="#FF8A65", width=2, dash="dash"), mode="lines+markers",
        marker=dict(size=4, symbol="triangle-up"),
        fill="tonexty", fillcolor="rgba(255,138,101,0.08)"))
    # Auto-ajustement : marge 20% au-dessus du max des rafales
    y_auto = max(df[g_col].max() * 1.2, 15) if not df[g_col].dropna().empty else 50
    fig.update_layout(
        title=dict(text=f"⚡ Vent à {height}", font=dict(size=12, color="white")),
        paper_bgcolor="#0E1117", plot_bgcolor="#161B2E",
        font=dict(color="white", size=9), height=260,
        margin=dict(l=50, r=50, t=35, b=55),
        legend=dict(orientation="h", y=-0.3, font=dict(size=8)),
        yaxis=dict(title="km/h", gridcolor="#2a2a3a", range=[0, y_auto]),
        xaxis=dict(gridcolor="#2a2a3a", tickangle=-45, tickfont=dict(size=7)),
        hovermode="x unified",
    )
    return fig

def bt_make_wind_rose(df, dir_col, spd_col, title=None):
    """Rose des vents pour Bénin Terminal (directions en FR, vitesses en km/h)."""
    if dir_col not in df.columns or spd_col not in df.columns:
        return go.Figure()
    DIR_MAP = {"N":0,"NNE":22.5,"NE":45,"ENE":67.5,"E":90,"ESE":112.5,
               "SE":135,"SSE":157.5,"S":180,"SSO":202.5,"SO":225,"OSO":247.5,
               "O":270,"ONO":292.5,"NO":315,"NNO":337.5,
               "SSW":202.5,"SW":225,"WSW":247.5,"W":270,"WNW":292.5,"NW":315,"NNW":337.5}
    dirs = df[dir_col].map(lambda x: DIR_MAP.get(str(x).strip().upper(), None)).dropna()
    spds = df[spd_col][dirs.index]
    if len(dirs) == 0:
        return go.Figure()
    bins = [0,10,20,30,40,100]
    labels = ["0-10","10-20","20-30","30-40",">40"]
    colors = ["#4FC3F7","#81C784","#FFB74D","#E67E22","#E74C3C"]
    dir_labels = ["N","NNE","NE","ENE","E","ESE","SE","SSE","S","SSO","SO","OSO","O","ONO","NO","NNO"]
    sector_size = 22.5
    traces = []
    for i, (lo, hi) in enumerate(zip(bins[:-1], bins[1:])):
        mask = (spds >= lo) & (spds < hi)
        sector_dirs = dirs[mask]
        r = []
        for s_deg in [j*sector_size for j in range(16)]:
            count = ((sector_dirs >= s_deg - sector_size/2) & (sector_dirs < s_deg + sector_size/2)).sum()
            r.append(count)
        traces.append(go.Barpolar(r=r, theta=dir_labels, name=f"{labels[i]} km/h",
            marker_color=colors[i], opacity=0.85))
    fig = go.Figure(traces)
    fig.update_layout(
        title=dict(text=title or "Rose des vents", font=dict(size=11, color="white")),
        paper_bgcolor="#0E1117", plot_bgcolor="#0E1117",
        font=dict(color="white", size=9), height=300,
        margin=dict(l=20, r=20, t=45, b=20),
        polar=dict(
            bgcolor="#161B2E",
            radialaxis=dict(showticklabels=False, gridcolor="#2a2a3a"),
            angularaxis=dict(direction="clockwise", gridcolor="#2a2a3a"),
        ),
        legend=dict(orientation="h", y=-0.12, font=dict(size=8)),
        showlegend=True,
    )
    return fig


def bt_generate_demo():
    """Données démo Bénin Terminal."""
    import numpy as np
    now = now_local()
    times = [now + timedelta(hours=2*i) for i in range(21)]
    wx_list = ["Assez nuageux","Modérément nuageux","Averses de pluies faibles"]
    return pd.DataFrame({
        "forecast_time_local": times,
        "V10m_Dir":    ["SO"]*21,
        "V10m_Km/h":   [10+5*np.sin(i/3) for i in range(21)],
        "RafaleV10_Km/h": [20+8*np.sin(i/3) for i in range(21)],
        "V22m_Dir":    ["SO"]*21,
        "V22m_Km/h":   [11+5*np.sin(i/3) for i in range(21)],
        "RafaleV22_Km/h": [22+9*np.sin(i/3) for i in range(21)],
        "V60m_Dir":    ["SO"]*21,
        "V60m_Km/h":   [13+6*np.sin(i/3) for i in range(21)],
        "RafaleV60_Km/h": [26+10*np.sin(i/3) for i in range(21)],
        "V70m_Dir":    ["SO"]*21,
        "V70m_Km/h":   [14+6*np.sin(i/3) for i in range(21)],
        "RafaleV70_Km/h": [28+11*np.sin(i/3) for i in range(21)],
        "T(°C)":       [27+np.sin(i/4) for i in range(21)],
        "Pluie(%)":    [10]*21,
        "Visibilite_km": [9]*21,
        "Temps_sensible": [wx_list[i%3] for i in range(21)],
    })


def read_bt_pdf_corrections(pdf_bytes):
    """
    Lit le PDF bulletin Bénin Terminal et extrait les corrections :
    T°C, Temps sensible, Probabilité de pluie, Visibilité.
    Retourne un DataFrame avec colonnes :
    forecast_time_local, T(°C), Temps_sensible, Pluie(%), Visibilite_km
    """
    try:
        import pdfplumber, io, re
        rows = []
        # Mapping temps sensible PDF → liste déroulante dashboard
        WX_MAP = {
            "ensoleillé":               "Ensoleillé",
            "peu nuageux":              "Peu nuageux",
            "modérément nuageux":       "Modérément nuageux",
            "moderement nuageux":       "Modérément nuageux",
            "assez nuageux":            "Assez nuageux",
            "couvert":                  "Couvert",
            "orages et pluies":         "Orages et pluies",
            "averses de fortes pluies": "Averses de fortes pluies",
            "averses de pluies faibles":"Averses de pluies faibles",
            "poussière":                "Poussière",
            "poussiere":                "Poussière",
            "brume sèche":              "Brume sèche",
            "brume seche":              "Brume sèche",
            "brouillard":               "Brouillard/Brume humide",
            "brume humide":             "Brouillard/Brume humide",
        }

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if not row or len(row) < 5:
                            continue
                        # Chercher lignes avec heure (ex: 21h, 01h, 07h...)
                        heure_cell = str(row[1] or "").strip()
                        date_cell  = str(row[0] or "").strip()
                        if not re.match(r"[0-9]{1,2}h", heure_cell):
                            continue
                        try:
                            # T°C — colonne 13 (index fixe dans le bulletin)
                            t_val = None
                            for ci in range(min(len(row), 16)):
                                cell = str(row[ci] or "").strip().replace(",",".")
                                try:
                                    v = float(cell)
                                    if 15 <= v <= 45:
                                        t_val = v
                                        break
                                except: pass

                            # Temps sensible — chercher cellule texte météo
                            wx_val = None
                            for ci in range(len(row)):
                                cell_low = str(row[ci] or "").strip().lower()
                                for key, mapped in WX_MAP.items():
                                    if key in cell_low:
                                        wx_val = mapped
                                        break
                                if wx_val: break

                            # Pluie % — chercher entier entre 0 et 100
                            pluie_val = None
                            for ci in range(len(row)):
                                cell = str(row[ci] or "").strip()
                                try:
                                    v = float(cell)
                                    if 0 <= v <= 100 and v != t_val:
                                        pluie_val = v
                                except: pass

                            # Visibilité — chercher valeur entre 1 et 20
                            vis_val = None
                            for ci in range(len(row)-1, -1, -1):
                                cell = str(row[ci] or "").strip()
                                try:
                                    v = float(cell)
                                    if 1 <= v <= 20:
                                        vis_val = v
                                        break
                                except: pass

                            if t_val or wx_val:
                                rows.append({
                                    "heure_str":    heure_cell,
                                    "date_str":     date_cell,
                                    "T(°C)":        t_val,
                                    "Temps_sensible": wx_val,
                                    "Pluie(%)":     pluie_val,
                                    "Visibilite_km": vis_val,
                                })
                        except Exception:
                            continue

        if not rows:
            return None, "Aucune donnée extraite du PDF"

        df_corr = pd.DataFrame(rows)
        return df_corr, None

    except Exception as e:
        return None, str(e)


# Pictogrammes temps sensible Bénin Terminal
BT_WX_ICONS = {
    "Ensoleillé":               "☀️",
    "Peu nuageux":              "🌤️",
    "Modérément nuageux":       "⛅",
    "Assez nuageux":            "🌥️",
    "Couvert":                  "☁️",
    "Orages et pluies":         "⛈️",
    "Averses de fortes pluies": "🌧️",
    "Averses de pluies faibles":"🌦️",
    "Poussière":                "🌫️",
    "Brume sèche":              "🌫️",
    "Brouillard/Brume humide":  "🌁",
}

BT_WX_LIST = list(BT_WX_ICONS.keys())


def render_benin_terminal():
    """Affichage complet Bénin Terminal — même architecture que Sème."""
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    import numpy as np

    # ── En-tête ──────────────────────────────────────────────────────────────
    now_loc = now_local()
    st.markdown(f"""
    <div class="marine-header">
        <div>
            <img src="{_LOGO_BT_URI}"
                 style='height:55px;object-fit:contain;vertical-align:middle;margin-right:1rem;'/>
        </div>
        <div>
            <div class="subtitle">METEO-BENIN · DPROM / SPAM</div>
            <h1>BULLETIN SPÉCIAL BÉNIN TERMINAL</h1>
            <div style="color:#adb5bd;font-size:0.78rem;margin-top:0.2rem;">
                📍 6.35°N, 2.43°E · Port de Cotonou, Bénin &nbsp;|&nbsp;
                Source : ECMWF via GEE &nbsp;|&nbsp;
                Mis à jour : {now_loc.strftime("%d/%m/%Y %H:%M")} (UTC+1)
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

    # ── Chargement CSV + Filtre temporel dans la sidebar ────────────────────
    is_demo   = False
    user_role = st.session_state.get("user_role", "client")

    with st.sidebar:
        st.divider()
        st.markdown("### ⚓ Bénin Terminal")

        # Upload CSV + PDF — admin uniquement
        if user_role == "admin":
            st.markdown("**📂 Bulletin CSV**")
            bt_file = st.file_uploader("Charger CSV", type=["csv"], key="bt_uploader")
            if bt_file:
                df_bt = pd.read_csv(bt_file)
                df_bt["forecast_time_local"] = pd.to_datetime(df_bt["forecast_time_local"])
                for col, default in [("T(°C)", 28.0), ("Pluie(%)", 10.0),
                                     ("Visibilite_km", 9.0), ("Temps_sensible","Assez nuageux")]:
                    if col not in df_bt.columns:
                        df_bt[col] = default
                st.session_state["bt_df"] = df_bt
                st.success(f"✅ {len(df_bt)} échéances chargées")

            st.markdown("**📄 PDF corrigé**")
            st.caption("T°C, Temps sensible, Pluie, Visibilité")
            bt_pdf = st.file_uploader("Charger PDF bulletin", type=["pdf"], key="bt_pdf_uploader")
            if bt_pdf and "bt_df" in st.session_state:
                pdf_bytes = bt_pdf.read()
                df_corr, err = read_bt_pdf_corrections(pdf_bytes)
                if err:
                    st.warning(f"⚠️ PDF : {err}")
                elif df_corr is not None:
                    # Appliquer corrections sur bt_df
                    df_base = st.session_state["bt_df"].copy()
                    n_corr = 0
                    for _, cr in df_corr.iterrows():
                        h_str = str(cr["heure_str"]).replace("h","").zfill(2)
                        mask  = df_base["forecast_time_local"].dt.strftime("%H") == h_str
                        if mask.any():
                            if cr["T(°C)"] is not None:
                                df_base.loc[mask, "T(°C)"] = cr["T(°C)"]
                            if cr["Temps_sensible"] is not None:
                                df_base.loc[mask, "Temps_sensible"] = cr["Temps_sensible"]
                            if cr["Pluie(%)"] is not None:
                                df_base.loc[mask, "Pluie(%)"] = cr["Pluie(%)"]
                            if cr["Visibilite_km"] is not None:
                                df_base.loc[mask, "Visibilite_km"] = cr["Visibilite_km"]
                            n_corr += 1
                    st.session_state["bt_df"] = df_base
                    st.success(f"✅ {n_corr} échéances corrigées depuis PDF")
            elif bt_pdf and "bt_df" not in st.session_state:
                st.warning("⚠️ Chargez d'abord le CSV avant le PDF")
            st.divider()

        # Chargement données
        if "bt_df" in st.session_state:
            df_bt = st.session_state["bt_df"]
        else:
            df_bt = bt_generate_demo()
            is_demo = True

        # Filtre temporel
        df_bt["forecast_time_local"] = pd.to_datetime(df_bt["forecast_time_local"])
        times  = sorted(df_bt["forecast_time_local"].dt.to_pydatetime().tolist())
        dt_min = times[0]; dt_max = times[-1]
        _19h   = dt_min.replace(hour=19, minute=0, second=0, microsecond=0)
        _19h_times = [t for t in times if t >= _19h]
        _def_start = _19h_times[0] if _19h_times else dt_min

        st.markdown("#### 🕐 Période")
        show_past = st.toggle("Voir les échéances passées", value=False, key="bt_show_past")
        def_start = dt_min if show_past else _def_start

        col_fa, col_fb = st.columns(2)
        with col_fa:
            s_d = st.date_input("De", value=def_start.date(),
                                min_value=dt_min.date(), max_value=dt_max.date(), key="bt_sd")
            _hs = sorted({t.hour for t in times if t.date()==s_d}) or list(range(0,24,2))
            _def_sh = def_start.hour if s_d == def_start.date() else _hs[0]
            _sh_idx = _hs.index(_def_sh) if _def_sh in _hs else 0
            s_h = st.selectbox("Heure", _hs, format_func=lambda h:f"{h:02d}:00",
                               index=_sh_idx, key="bt_sh")
        with col_fb:
            e_d = st.date_input("À", value=dt_max.date(),
                                min_value=dt_min.date(), max_value=dt_max.date(), key="bt_ed")
            _he = sorted({t.hour for t in times if t.date()==e_d}) or list(range(0,24,2))
            e_h = st.selectbox("Heure", _he, format_func=lambda h:f"{h:02d}:00",
                               index=len(_he)-1, key="bt_eh")

        if is_demo:
            st.caption("⚠️ Mode démonstration")

    # Badge démo
    demo_badge = " · <span style='color:#ffa94d;font-size:0.7rem;'>DEMO</span>" if is_demo else ""

    t_start = datetime.combine(s_d, datetime.min.time()).replace(hour=s_h)
    t_end   = datetime.combine(e_d, datetime.min.time()).replace(hour=e_h)
    df_f = df_bt[(df_bt["forecast_time_local"] >= pd.Timestamp(t_start)) &
                 (df_bt["forecast_time_local"] <= pd.Timestamp(t_end))].copy()
    if df_f.empty: df_f = df_bt.copy()

    # ── Alerte globale ────────────────────────────────────────────────────────
    lbl_a, css_a, msg_a = bt_global_alert(df_f)
    st.markdown(f'''<div class="warning-box {css_a}"><b>{lbl_a}</b> — {msg_a}</div>''',
                unsafe_allow_html=True)

    # ── KPI ───────────────────────────────────────────────────────────────────
    now = now_local()
    closest_idx = (df_f["forecast_time_local"] - now).abs().argsort().iloc[0]
    row = df_f.iloc[closest_idx]
    colors_map = {"green":"#2ECC71","yellow":"#F1C40F","orange":"#E67E22","red":"#E74C3C"}

    k1,k2,k3,k4,k5,k6 = st.columns(6)
    for col_w, height, g_col, label in [
        (k1,"10m","RafaleV10_Km/h","Rafales 10m"),
        (k2,"22m","RafaleV22_Km/h","Rafales 22m"),
        (k3,"60m","RafaleV60_Km/h","Rafales 60m"),
        (k4,"70m","RafaleV70_Km/h","Rafales 70m"),
    ]:
        v = float(row.get(g_col, 0) or 0)
        lvl, lbl2 = bt_get_alert(v, height)
        with col_w:
            st.markdown(f"""
            <div class='kpi-card' style='border-top:3px solid {colors_map[lvl]};text-align:center;'>
                <div class='kpi-label'>{label}</div>
                <div class='kpi-value'>{v:.0f}</div>
                <div class='kpi-unit'>km/h · {lbl2}</div>
            </div>""", unsafe_allow_html=True)
    with k5:
        t_val = float(row.get("T(°C)", 0) or 0)
        st.markdown(f"""
        <div class='kpi-card' style='text-align:center;'>
            <div class='kpi-label'>🌡️ T°C</div>
            <div class='kpi-value'>{t_val:.1f}</div>
            <div class='kpi-unit'>°C</div>
        </div>""", unsafe_allow_html=True)
    with k6:
        wx = str(row.get("Temps_sensible","—"))
        icon = BT_WX_ICONS.get(wx, "🌤️")
        st.markdown(f"""
        <div class='kpi-card' style='text-align:center;'>
            <div class='kpi-label'>Temps</div>
            <div class='kpi-value'>{icon}</div>
            <div class='kpi-unit' style='font-size:0.65rem;'>{wx}</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("---")

    # ── Onglets ───────────────────────────────────────────────────────────────
    tab1, tab2, tab3, tab4 = st.tabs(["🌤️ Météo", "💨 Vent", "📋 Données", "💾 Export"])

    # ── Onglet Météo ─────────────────────────────────────────────────────────
    with tab1:
        # Labels X exacts depuis le CSV
        x = df_f["forecast_time_local"].dt.strftime("%a %d/%m %Hh")
        pluie_col = "Pluie(%)" if "Pluie(%)" in df_f.columns else "Pluie(mm)"

        # ── Graphique T°C seul + pictogrammes au-dessus ──────────────────
        fig_meteo = go.Figure()

        # Courbe température
        fig_meteo.add_trace(go.Scatter(
            x=x, y=df_f["T(°C)"],
            name="T°C",
            line=dict(color="#FF6B6B", width=2.5),
            mode="lines+markers",
            marker=dict(size=6, color="#FF6B6B"),
            fill="tozeroy", fillcolor="rgba(255,107,107,0.08)",
        ))

        # Pictogrammes au-dessus de la courbe
        if "Temps_sensible" in df_f.columns:
            t_vals = df_f["T(°C)"].fillna(0)
            t_max  = t_vals.max() if t_vals.max() > 0 else 35
            for i, (_, r) in enumerate(df_f.iterrows()):
                wx_val = str(r.get("Temps_sensible", ""))
                icon   = BT_WX_ICONS.get(wx_val, "🌤️")
                t_val  = float(r.get("T(°C)", 0) or 0)
                fig_meteo.add_annotation(
                    x=x.iloc[i] if hasattr(x, "iloc") else x[i],
                    y=t_val,
                    text=icon,
                    showarrow=False,
                    font=dict(size=16),
                    yanchor="bottom",
                    yshift=8,
                )

        fig_meteo.update_layout(
            title=dict(text="🌡️ Température & Temps sensible",
                       font=dict(size=12, color="white")),
            paper_bgcolor="#0E1117", plot_bgcolor="#161B2E",
            font=dict(color="white", size=9), height=320,
            margin=dict(l=55, r=30, t=50, b=70),
            yaxis=dict(
                title="T (°C)", gridcolor="#2a2a3a", color="#FF6B6B",
                range=[20, 35], dtick=2,
                tickmode="linear", tick0=20,
            ),
            xaxis=dict(gridcolor="#2a2a3a", tickangle=-45, tickfont=dict(size=8)),
            hovermode="x unified",
            showlegend=False,
        )
        st.plotly_chart(fig_meteo, use_container_width=True)

        # ── Pluie + Visibilité combinées ─────────────────────────────────
        if "Visibilite_km" in df_f.columns:
            fig_pv = make_subplots(specs=[[{"secondary_y": True}]])

            # Barres pluie (axe gauche)
            fig_pv.add_trace(go.Bar(
                x=x, y=df_f[pluie_col],
                name="Pluie (%)",
                marker_color="rgba(100,181,246,0.6)",
                marker_line_width=0,
            ), secondary_y=False)

            # Courbe visibilité (axe droit)
            fig_pv.add_trace(go.Scatter(
                x=x, y=df_f["Visibilite_km"],
                name="Visibilité (km)",
                line=dict(color="#A9CCE3", width=2.5),
                mode="lines+markers", marker=dict(size=6),
            ), secondary_y=True)

            fig_pv.update_layout(
                title=dict(text="🌧️ Probabilité de pluie & 👁️ Visibilité minimale",
                           font=dict(size=12, color="white")),
                paper_bgcolor="#0E1117", plot_bgcolor="#161B2E",
                font=dict(color="white", size=9), height=260,
                margin=dict(l=55, r=65, t=45, b=65),
                legend=dict(orientation="h", y=-0.28, font=dict(size=9)),
                xaxis=dict(gridcolor="#2a2a3a", tickangle=-45, tickfont=dict(size=8)),
                hovermode="x unified", bargap=0.15,
            )
            fig_pv.update_yaxes(
                title_text="Pluie (%)", secondary_y=False,
                gridcolor="#2a2a3a", color="#64B5F6",
                range=[0, (df_f[pluie_col].max() or 100) * 1.3],
            )
            fig_pv.update_yaxes(
                title_text="Visibilité (km)", secondary_y=True,
                color="#A9CCE3", gridcolor="rgba(0,0,0,0)",
            )
            st.plotly_chart(fig_pv, use_container_width=True)

    # ── Onglet Vent ──────────────────────────────────────────────────────────
    with tab2:
        # Chaque niveau : graphique (gauche 70%) + rose (droite 30%)
        levels = [
            ("10m", "V10m_Dir", "V10m_Km/h", "RafaleV10_Km/h"),
            ("22m", "V22m_Dir", "V22m_Km/h", "RafaleV22_Km/h"),
            ("60m", "V60m_Dir", "V60m_Km/h", "RafaleV60_Km/h"),
            ("70m", "V70m_Dir", "V70m_Km/h", "RafaleV70_Km/h"),
        ]
        for height, dir_col, v_col, g_col in levels:
            col_graph, col_rose = st.columns([3, 1])
            with col_graph:
                # Labels X exacts
                x_w = df_f["forecast_time_local"].dt.strftime("%a %d/%m %Hh")
                fig_w = bt_plot_wind(df_f, height, v_col, g_col)
                # Mettre à jour les labels X dans la figure
                for trace in fig_w.data:
                    trace.x = list(x_w)
                st.plotly_chart(fig_w, use_container_width=True)
            with col_rose:
                st.plotly_chart(
                    bt_make_wind_rose(df_f, dir_col, v_col, f"Rose {height}"),
                    use_container_width=True
                )

        # Légende seuils
        st.markdown("#### 📌 Seuils d'alerte rafales (km/h)")
        sl1, sl2 = st.columns(2)
        with sl1:
            st.markdown("| Hauteur | 🟢 Vert | 🟡 Jaune | 🟠 Orange | 🔴 Rouge |\n|--|--|--|--|--|\n| 10m | ≤28 | 29–49 | 50–74 | ≥75 |\n| 22m | ≤33 | 34–57 | 58–86 | ≥87 |")
        with sl2:
            st.markdown("| Hauteur | 🟢 Vert | 🟡 Jaune | 🟠 Orange | 🔴 Rouge |\n|--|--|--|--|--|\n| 60m | ≤40 | 41–70 | 71–107 | ≥107 |\n| 70m | ≤40 | 42–72 | 73–109 | ≥110 |")

    # ── Onglet Données ───────────────────────────────────────────────────────
    with tab3:
        st.markdown("### 📋 Tableau complet")
        disp = df_f.copy()
        disp["forecast_time_local"] = disp["forecast_time_local"].dt.strftime("%d/%m %H:%M")
        for col in ["V10m_Km/h","RafaleV10_Km/h","V22m_Km/h","RafaleV22_Km/h",
                    "V60m_Km/h","RafaleV60_Km/h","V70m_Km/h","RafaleV70_Km/h"]:
            if col in disp.columns: disp[col] = disp[col].round(1)
        if "T(°C)" in disp.columns: disp["T(°C)"] = disp["T(°C)"].round(1)

        def style_bt(row):
            styles = [""] * len(row)
            cols_d = list(disp.columns)
            cmap = {"green":"background-color:#1a3a2a;color:#2ECC71",
                    "yellow":"background-color:#3a3a1a;color:#F1C40F",
                    "orange":"background-color:#3a2a1a;color:#E67E22",
                    "red":"background-color:#3a1a1a;color:#E74C3C"}
            for col_name, height in [("RafaleV10_Km/h","10m"),("RafaleV22_Km/h","22m"),
                                      ("RafaleV60_Km/h","60m"),("RafaleV70_Km/h","70m")]:
                if col_name in cols_d:
                    idx = cols_d.index(col_name)
                    lvl, _ = bt_get_alert(float(row[col_name] or 0), height)
                    styles[idx] = cmap[lvl]
            return styles

        st.dataframe(disp.style.apply(style_bt, axis=1),
                     use_container_width=True, height=420)

    # ── Onglet Export ────────────────────────────────────────────────────────
    with tab4:
        st.markdown("### 💾 Export des données")
        exp = df_f.copy()
        exp["forecast_time_local"] = exp["forecast_time_local"].dt.strftime("%Y-%m-%d %H:%M")
        fname = f"BeninTerminal_{now_local().strftime('%d%m%Y_%H%M')}.csv"
        st.download_button("⬇️ Télécharger CSV filtré",
            data=exp.to_csv(index=False).encode("utf-8"),
            file_name=fname, mime="text/csv")

    # Footer
    st.markdown("""
    <div style='text-align:center;color:#4a6480;font-size:0.68rem;margin-top:2rem;
                padding:1rem 0;border-top:1px solid rgba(46,117,182,0.15);'>
        © 2026 · LAOUROU MAKONDJOU DIANE · Météorologiste & Data Scientist · METEO-BENIN / DPROM / SPAM<br>
        Source : ECMWF Open Data via Google Earth Engine
    </div>""", unsafe_allow_html=True)


if __name__ == "__main__":
    main()