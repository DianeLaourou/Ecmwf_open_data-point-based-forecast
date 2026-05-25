# word_exporter.py — Bulletin Marine Word complet
import io, math
import pandas as pd
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import config

# =============================================================================
# ECMWF OpenCharts — téléchargement automatique des images
# =============================================================================

ECMWF_BASE = "https://charts.ecmwf.int/opencharts-api/v1/products"

def _ecmwf_base_time(run_datetime: datetime) -> str:
    """Formate run_datetime en base_time ISO8601 pour l'API ECMWF."""
    # On utilise toujours 00Z ou 12Z — aligner sur la run 00Z du jour du run
    from datetime import datetime as _dt
    today_00z = _dt.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
    run_00z   = run_datetime.replace(hour=0, minute=0, second=0, microsecond=0)
    base = today_00z if today_00z >= run_00z else run_00z
    return base.strftime("%Y-%m-%dT%H:%M:%SZ")

def _ecmwf_valid_time(run_datetime: datetime) -> str:
    """J+1 à 12UTC = valid_time pour Figure 3."""
    vt = run_datetime.replace(hour=0, minute=0, second=0, microsecond=0) + timedelta(days=1, hours=12)
    return vt.strftime("%Y-%m-%dT%H:%M:%SZ")

def fetch_ecmwf_image(url: str, label: str = "") -> io.BytesIO | None:
    """
    Télécharge une image depuis l'API ECMWF OpenCharts.
    L'API retourne d'abord un JSON contenant data.link.href → URL de l'image PNG.
    Retourne un BytesIO prêt à être inséré dans python-docx, ou None si échec.
    """
    try:
        import requests
        headers = {"User-Agent": "Mozilla/5.0 (METEO-BENIN pipeline)"}

        # Étape 1 : appel API → JSON avec l'URL de l'image
        resp = requests.get(url, headers=headers, timeout=30)
        resp.raise_for_status()

        content_type = resp.headers.get("Content-Type", "")

        if "image" in content_type:
            # Cas direct (rare) — déjà une image
            img_url = None
            img_data = resp.content
        elif "json" in content_type or "text" in content_type:
            # Cas normal — extraire l'URL depuis le JSON
            data = resp.json()
            img_url = data.get("data", {}).get("link", {}).get("href")
            if not img_url:
                print(f"  ⚠️  {label} : champ data.link.href absent dans la réponse JSON")
                return None
            # Étape 2 : téléchargement de l'image
            resp2 = requests.get(img_url, headers=headers, timeout=60)
            resp2.raise_for_status()
            img_data = resp2.content
        else:
            print(f"  ⚠️  {label} : Content-Type inattendu ({content_type})")
            return None

        buf = io.BytesIO(img_data)
        buf.seek(0)
        print(f"  ✅ Image téléchargée : {label}")
        return buf

    except Exception as e:
        print(f"  ⚠️  Échec téléchargement {label} : {e}")
        return None

def build_ecmwf_urls(run_datetime: datetime) -> dict:
    """
    Construit les 5 URLs ECMWF en fonction du run_datetime.
    Retourne un dict avec les clés : ensgram, swh_dir, swh_prob, mwp_8s, mwp_15s
    """
    from urllib.parse import quote
    bt = _ecmwf_base_time(run_datetime)   # base_time = 00Z du jour du run
    vt = _ecmwf_valid_time(run_datetime)  # valid_time = J+1 12UTC
    proj = "opencharts_northern_africa"

    def enc(s): return quote(s, safe="")

    # station_name : encoder en ASCII pur (enlever accents)
    import unicodedata
    station_name = unicodedata.normalize("NFKD", config.POINT["name"])
    station_name = station_name.encode("ascii", "ignore").decode("ascii")

    return {
        "ensgram": (
            f"{ECMWF_BASE}/opencharts_meteogram/"
            f"?base_time={enc(bt)}"
            f"&epsgram=classical_wave"
            f"&lat={config.POINT['lat']}"
            f"&lon={config.POINT['lon']}"
            f"&station_name={enc(station_name)}"
        ),
        "swh_dir": (
            f"{ECMWF_BASE}/medium-swh-mwd/"
            f"?base_time={enc(bt)}"
            f"&valid_time={enc(vt)}"
            f"&projection={proj}"
        ),
        "swh_prob": (
            f"{ECMWF_BASE}/medium-swh-probability/"
            f"?base_time={enc(bt)}"
            f"&valid_time={enc(vt)}"
            f"&projection={proj}"
        ),
        "mwp_8s": (
            f"{ECMWF_BASE}/medium-mwp-probability/"
            f"?base_time={enc(bt)}"
            f"&valid_time={enc(vt)}"
            f"&projection={proj}"
        ),
        "mwp_15s": (
            f"{ECMWF_BASE}/medium-mwp-probability/"
            f"?base_time={enc(bt)}"
            f"&valid_time={enc(vt)}"
            f"&projection={proj}"
            f"&param=mwpg15"
        ),
    }

LOGO_REP = Path("D:/pipeline/logo_republique.png")
LOGO_MET = Path("D:/pipeline/logo_meteo_oval.png")   # logo oval seul (col droite en-tête)
BLUE_MED = RGBColor(0x2E,0x75,0xB6); WHITE = RGBColor(0xFF,0xFF,0xFF); BLACK = RGBColor(0,0,0)
ENG_DS = {0:"Mon.",1:"Tue.",2:"Wed.",3:"Thu.",4:"Fri.",5:"Sat.",6:"Sun."}
ENG_MS = {1:"Jan",2:"Feb",3:"Mar",4:"Apr",5:"May",6:"Jun",7:"Jul",8:"Aug",9:"Sep",10:"Oct",11:"Nov",12:"Dec"}
ENG_DF = {0:"Monday",1:"Tuesday",2:"Wednesday",3:"Thursday",4:"Friday",5:"Saturday",6:"Sunday"}
ENG_MF = {1:"January",2:"February",3:"March",4:"April",5:"May",6:"June",7:"July",8:"August",9:"September",10:"October",11:"November",12:"December"}
DAY_BG = {0:"E2EFDA",1:"FFFFFF",2:"E2EFDA",3:"FFFFFF",4:"E2EFDA",5:"FFFFFF",6:"E2EFDA"}
WX_LIST = [
    "Sunny", "Mostly sunny", "Partly cloudy", "Mostly cloudy", "Cloudy", "Overcast",
    "Light rain", "Moderate rain", "Heavy rain", "Rain showers",
    "Thunderstorm", "Squally",
    "Mist", "Fog", "Haze",
]

# Nombre de pas de temps par jour (8 pas de 3h sur 24h)
STEPS_PER_DAY = 8

def ordinal(n):
    if 11<=n<=13: return f"{n}th"
    return f"{n}"+{1:"st",2:"nd",3:"rd"}.get(n%10,"th")

def fmt_period(dt):
    return f"{ENG_DF[dt.weekday()]}, {ordinal(dt.day)} {ENG_MF[dt.month]} {dt.year} at 07:00 pm"

def fmt(v,d=1):
    if v is None or (isinstance(v,float) and math.isnan(v)): return "—"
    return f"{v:.{d}f}" if d>0 else str(int(round(v)))

def sbg(cell,col):
    tc=cell._tc; p=tc.get_or_add_tcPr()
    for o in p.findall(qn('w:shd')): p.remove(o)
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),col); p.append(s)

def smg(cell,t=20,b=20,l=25,r=25):
    tc=cell._tc; p=tc.get_or_add_tcPr(); m=OxmlElement('w:tcMar')
    for sd,v in [('top',t),('bottom',b),('left',l),('right',r)]:
        el=OxmlElement(f'w:{sd}'); el.set(qn('w:w'),str(v)); el.set(qn('w:type'),'dxa'); m.append(el)
    p.append(m)

def sbo(cell):
    tc=cell._tc; p=tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for sd in ['top','bottom','left','right']:
        el=OxmlElement(f'w:{sd}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); el.set(qn('w:color'),'BBBBBB'); b.append(el)
    p.append(b)

def snb(cell):
    tc=cell._tc; p=tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for sd in ['top','bottom','left','right','insideH','insideV']:
        el=OxmlElement(f'w:{sd}'); el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0'); el.set(qn('w:color'),'FFFFFF'); b.append(el)
    p.append(b)

def cp(cell,text,sz=7,bold=False,color=None,italic=False):
    p=cell.paragraphs[0]; p.clear(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(str(text)); r.font.size=Pt(sz); r.font.bold=bold; r.font.italic=italic
    if color: r.font.color.rgb=color
    return p

def ar(para,text,sz=9,bold=False,color=None,italic=False):
    r=para.add_run(str(text)); r.font.size=Pt(sz); r.font.bold=bold; r.font.italic=italic
    if color: r.font.color.rgb=color
    return r

def ybar(doc):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    pp=p._p.get_or_add_pPr(); s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),'FFD700'); pp.append(s)
    sp=OxmlElement('w:spacing'); sp.set(qn('w:before'),'0'); sp.set(qn('w:after'),'0'); sp.set(qn('w:line'),'100'); sp.set(qn('w:lineRule'),'exact'); pp.append(sp)
    r=p.add_run(" "); r.font.size=Pt(5)
    rp=OxmlElement('w:rPr'); rs=OxmlElement('w:shd'); rs.set(qn('w:val'),'clear'); rs.set(qn('w:color'),'auto'); rs.set(qn('w:fill'),'FFD700'); rp.append(rs); r._r.append(rp)

def tzone(doc,ph,sz=9):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    pp=p._p.get_or_add_pPr(); s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),'FFFDE7'); pp.append(s)
    r=p.add_run(ph); r.font.size=Pt(sz); r.font.italic=True; r.font.color.rgb=RGBColor(0x99,0x99,0x99)

def izone(doc,ph,hcm=10):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
    pp=p._p.get_or_add_pPr(); s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),'F0F8FF'); pp.append(s)
    sp=OxmlElement('w:spacing'); sp.set(qn('w:before'),str(int(hcm*567))); sp.set(qn('w:after'),str(int(hcm*567))); pp.append(sp)
    r=p.add_run(ph); r.font.size=Pt(10); r.font.italic=True; r.font.color.rgb=RGBColor(0xAA,0xAA,0xAA)

def add_dropdown(cell):
    """Liste déroulante Weather condition — 7pt fixe."""
    sdt=OxmlElement('w:sdt')
    sp=OxmlElement('w:sdtPr')
    a=OxmlElement('w:alias'); a.set(qn('w:val'),'Weather Condition'); sp.append(a)
    t=OxmlElement('w:tag'); t.set(qn('w:val'),'wx'); sp.append(t)
    # Forcer la police via rPr dans sdtPr
    rpr=OxmlElement('w:rPr')
    sz_pr=OxmlElement('w:sz'); sz_pr.set(qn('w:val'),'14')   # 7pt = 14 demi-points
    szcs=OxmlElement('w:szCs'); szcs.set(qn('w:val'),'14')
    rpr.append(sz_pr); rpr.append(szcs); sp.append(rpr)
    dd=OxmlElement('w:dropDownList')
    i0=OxmlElement('w:listItem'); i0.set(qn('w:displayText'),'— Select —'); i0.set(qn('w:value'),''); dd.append(i0)
    for w in WX_LIST:
        it=OxmlElement('w:listItem'); it.set(qn('w:displayText'),w); it.set(qn('w:value'),w); dd.append(it)
    sp.append(dd); sdt.append(sp)
    sc=OxmlElement('w:sdtContent')
    p=OxmlElement('w:p'); pp2=OxmlElement('w:pPr'); jc=OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); pp2.append(jc); p.append(pp2)
    r=OxmlElement('w:r')
    rp=OxmlElement('w:rPr')
    sz2=OxmlElement('w:sz'); sz2.set(qn('w:val'),'14')
    szcs2=OxmlElement('w:szCs'); szcs2.set(qn('w:val'),'14')
    rp.append(sz2); rp.append(szcs2); r.append(rp)
    tx=OxmlElement('w:t'); tx.text='— Select —'; r.append(tx); p.append(r); sc.append(p); sdt.append(sc)
    tc=cell._tc
    for ch in list(tc): tc.remove(ch)
    tc.append(sdt)


def add_confidence_dropdown(cell):
    """
    Liste déroulante Confidence : Low / Medium / High — 7pt.
    Préserve la fusion de cellule existante (ne supprime pas le vMerge).
    """
    tc = cell._tc
    # Supprimer uniquement les paragraphes existants, pas les attributs de fusion
    for child in list(tc):
        if child.tag in [qn('w:p'), qn('w:sdt')]:
            tc.remove(child)

    sdt = OxmlElement('w:sdt')
    sp  = OxmlElement('w:sdtPr')
    a   = OxmlElement('w:alias'); a.set(qn('w:val'),'Confidence'); sp.append(a)
    t   = OxmlElement('w:tag');   t.set(qn('w:val'),'conf');       sp.append(t)
    rpr = OxmlElement('w:rPr')
    sz_pr = OxmlElement('w:sz');   sz_pr.set(qn('w:val'),'14'); rpr.append(sz_pr)
    szcs  = OxmlElement('w:szCs'); szcs.set(qn('w:val'),'14');  rpr.append(szcs)
    sp.append(rpr)
    dd = OxmlElement('w:dropDownList')
    i0 = OxmlElement('w:listItem'); i0.set(qn('w:displayText'),'— Select —'); i0.set(qn('w:value'),''); dd.append(i0)
    for val in ["Low","Medium","High"]:
        it = OxmlElement('w:listItem'); it.set(qn('w:displayText'),val); it.set(qn('w:value'),val); dd.append(it)
    sp.append(dd); sdt.append(sp)
    sc = OxmlElement('w:sdtContent')
    p  = OxmlElement('w:p')
    pp2 = OxmlElement('w:pPr')
    jc  = OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); pp2.append(jc)
    p.append(pp2)
    r  = OxmlElement('w:r')
    rp = OxmlElement('w:rPr')
    sz2  = OxmlElement('w:sz');   sz2.set(qn('w:val'),'14'); rp.append(sz2)
    szcs2= OxmlElement('w:szCs'); szcs2.set(qn('w:val'),'14'); rp.append(szcs2)
    r.append(rp)
    tx = OxmlElement('w:t'); tx.text = '— Select —'; r.append(tx)
    p.append(r); sc.append(p); sdt.append(sc)
    tc.append(sdt)

def add_header(doc):
    # Tableau principal 3 colonnes
    t=doc.add_table(rows=1,cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,c in enumerate(t.rows[0].cells):
        c.width=[Cm(5.5),Cm(15),Cm(5)][i]; snb(c)

    # Col 0 — Logo République
    cl=t.rows[0].cells[0]; cl.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    p=cl.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_REP.exists(): p.add_run().add_picture(str(LOGO_REP),width=Cm(5))
    else: p.add_run("[Logo Rép.]")

    # Col 1 — Titre central
    cc=t.rows[0].cells[1]; cc.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    for txt,sz,bd in [
        ("DIRECTION DE LA PREVISION ET DU RESEAU D'OBSERVATION METEOROLOGIQUE",8.5,True),
        ("SERVICE DE LA PREVISION ET DE L'ASSISTANCE METEOROLOGIQUE",8.5,True),
        ("MARINE FORECAST FOR SEME",12,True)
    ]:
        p=cc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
        r=p.add_run(txt); r.bold=bd; r.font.size=Pt(sz)
    cc.paragraphs[0]._p.getparent().remove(cc.paragraphs[0]._p)

    # Col 2 — Logo Météo Bénin avec coordonnées intégrées
    cr=t.rows[0].cells[2]; cr.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    p=cr.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_MET.exists():
        p.add_run().add_picture(str(LOGO_MET), width=Cm(4.8))
    else:
        p.add_run("[Logo Météo Bénin]")

def make_full_table(doc, df_sorted):
    """
    Tableau complet toutes colonnes — 3 niveaux d'en-têtes.
    - Colonne Date (col 0) : fusionnée par jour + 2 cellules vides en en-tête
    - Colonne Confidence (col 22) : fusionnée par jour
    - Pagination : jours 1-2 sur page 1, jours 3-4 sur page 2
    """
    COLS = [
        ("Date",1.7),("Time",0.85),
        ("Dir.",0.65),("Spd\n(kts)",0.65),("Gust\n(kts)",0.65),
        ("Dir.",0.65),("Spd\n(kts)",0.65),
        ("MSLP\n(hPa)",0.8),("Vis.\n(km)",0.6),("T\n(°C)",0.6),
        ("Weather\ncond.",1.9),("Rain\n(%)",0.6),("SST\n(°C)",0.6),
        ("Dir.",0.6),("Per.\n(s)",0.6),("Ht.\n(m)",0.6),
        ("Dir.",0.6),("Per.\n(s)",0.6),("Ht.\n(m)",0.6),
        ("S.W\n(m)",0.65),
        ("Dir.",0.6),("Spd.\n(m/s)",0.65),
        ("Conf.",1.5),
    ]
    NC = len(COLS)
    GRP = [
        (0,1,"Date & Time"),(2,4,"Wind at 10m"),(5,6,"Wind at 100m"),
        (7,7,"MSLP (hPa)"),(8,8,"Vis. (km)"),(9,9,"T (°C)"),
        (10,10,"Weather condition"),(11,11,"Rain (%)"),(12,12,"SST (°C)"),
        (13,15,"Swell 1"),(16,18,"Swell 2"),(19,19,"S.W (m)"),
        (20,21,"Currents"),(22,22,"Confidence")
    ]

    # Regrouper les lignes par jour
    df_sorted = df_sorted.copy()
    df_sorted["_dt"] = pd.to_datetime(df_sorted["valid_local"])
    df_sorted["_day"] = df_sorted["_dt"].dt.date
    days = list(df_sorted["_day"].unique())

    # Nombre total de lignes de données
    N = len(df_sorted)

    # Créer le tableau : 3 lignes d'en-têtes + N lignes données
    tbl = doc.add_table(rows=3 + N, cols=NC)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── Ligne 0 : Weather / Ocean parameters ──────────────────────────────
    r0 = tbl.rows[0]
    # Col 0-1 (Date & Time) : vides mais bleu
    for j in range(2):
        sbg(r0.cells[j], "FFFFFF"); sbo(r0.cells[j])
    cw = r0.cells[2].merge(r0.cells[12])
    sbg(cw,"FFFFFF"); sbo(cw)
    cp(cw,"Weather parameters",sz=9,bold=True,color=RGBColor(0,0,0))
    cw.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    co = r0.cells[13].merge(r0.cells[22])
    sbg(co,"FFFFFF"); sbo(co)
    cp(co,"Ocean parameters",sz=9,bold=True,color=RGBColor(0,0,0))
    co.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ── Ligne 1 : groupes ─────────────────────────────────────────────────
    r1 = tbl.rows[1]; done = set()
    for gs, ge, lbl in GRP:
        if gs in done: continue
        c = r1.cells[gs].merge(r1.cells[ge]) if gs != ge else r1.cells[gs]
        sbg(c,"FFFFFF"); sbo(c); smg(c,15,15,20,20)
        cp(c, lbl, sz=7, bold=True, color=RGBColor(0,0,0))
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for i in range(gs, ge+1): done.add(i)

    # ── Ligne 2 : sous-en-têtes ───────────────────────────────────────────
    # Col 0 (Date) : afficher "Date & Time" sur cette ligne seulement
    r2 = tbl.rows[2]
    for j, (lbl, w) in enumerate(COLS):
        c = r2.cells[j]
        sbg(c,"FFFFFF"); sbo(c); smg(c,15,15,20,20)
        # Col 0 : "Date & Time" centré verticalement sur 3 lignes
        display_lbl = lbl if j != 0 else "Date & Time"
        cp(c, display_lbl, sz=6.5, bold=True,
           color=RGBColor(0,0,0))
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.width = Cm(w)

    # ── Répétition des en-têtes sur chaque page (tblHeader) ─────────────
    for row in tbl.rows[:3]:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        # Supprimer si déjà présent
        for existing in trPr.findall(qn('w:tblHeader')):
            trPr.remove(existing)
        # Insérer en première position
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), 'true')
        trPr.insert(0, tblHeader)

    # ── Données : par jour avec fusions ───────────────────────────────────
    row_offset = 0  # décalage dans le tableau (après 3 lignes d'en-têtes)

    for day_idx, day in enumerate(days):
        day_rows = df_sorted[df_sorted["_day"] == day]
        n_day = len(day_rows)
        # couleur alternée par ligne (voir row_bg ci-dessous)

        # Saut de page après le 2ème jour (entre jour 2 et jour 3)
        if day_idx == 2 and row_offset > 0:
            # Ajouter saut de page sur la première ligne du 3ème jour
            first_data_row = tbl.rows[3 + row_offset]
            first_cell = first_data_row.cells[1]
            fc_p = first_cell.paragraphs[0]._p
            fc_pPr = fc_p.get_or_add_pPr()
            pb = OxmlElement('w:pageBreakBefore')
            pb.set(qn('w:val'), '1')
            fc_pPr.append(pb)

        for k, (_, rd) in enumerate(day_rows.iterrows()):
            dt = pd.to_datetime(rd.get("valid_local"))
            swh = rd.get("swh_m")
            ws  = rd.get("wind10_spd_kt")
            i   = row_offset + k
            tr  = tbl.rows[3 + i]

            # Valeurs pour chaque colonne
            vals = [
                "",  # Date — sera fusionnée (col 0)
                dt.strftime("%H:%M"),  # Time (24h)
                str(rd.get("wind10_dir") or "—"), fmt(ws,0),
                fmt(rd.get("wind10_gust_kt"),0),
                str(rd.get("wind100_dir") or "—"),
                fmt(rd.get("wind100_spd_kt"),0),
                fmt(rd.get("mslp_hpa"),0), fmt(rd.get("vis_km"),0),
                fmt(rd.get("t2m_c"),0),
                None,  # Weather condition — liste déroulante (col 10)
                str(rd.get("rain_pct","—")), fmt(rd.get("sst_c"),0),
                str(rd.get("sw1_dir") or "—"),
                fmt(rd.get("sw1_period_s"),0), fmt(rd.get("sw1_ht_m"),1),
                str(rd.get("sw2_dir") or "—"),
                fmt(rd.get("sw2_period_s"),0), fmt(rd.get("sw2_ht_m"),1),
                fmt(swh,1),
                str(rd.get("cur_dir") or "—"),
                fmt(rd.get("cur_spd_ms"),2),
                "",  # Confidence — sera fusionnée (col 22)
            ]

            row_bg = "E2EFDA" if i % 2 == 0 else "FFFFFF"
            for j, v in enumerate(vals):
                c = tr.cells[j]; cb = row_bg
                # Pas de surlignage couleur sur les cellules données
                pass

                sbg(c,cb); sbo(c); smg(c,12,12,18,18)
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                if j == 10:   # Weather condition
                    sbg(c,"FFFACD"); add_dropdown(c)
                elif j == 22: # Confidence — liste déroulante
                    sbg(c,"E8F5E9"); add_confidence_dropdown(c)
                elif j == 0:  # Date — vide (sera fusionnée)
                    cp(c,"",sz=9)
                else:
                    fc = WHITE if cb=="FF0000" else None
                    cp(c, str(v) if v is not None else "—", sz=9, color=fc)

        # ── Fusionner colonne Date (col 0) sur tout le jour ───────────────
        if n_day > 1:
            first_row_idx = 3 + row_offset
            last_row_idx  = 3 + row_offset + n_day - 1
            merged_date = tbl.rows[first_row_idx].cells[0].merge(
                          tbl.rows[last_row_idx].cells[0])
            day_label = f"{ENG_DS[list(day_rows['_dt'])[0].weekday()]} {list(day_rows['_dt'])[0].day:02d} {ENG_MS[list(day_rows['_dt'])[0].month]}"
            sbg(merged_date, "E2EFDA" if row_offset % 2 == 0 else "FFFFFF")
            sbo(merged_date)
            smg(merged_date,12,12,18,18)
            cp(merged_date, day_label, sz=7, bold=True)
            merged_date.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # ── Fusionner colonne Confidence (col 22) sur tout le jour ────────
        if n_day > 1:
            first_row_idx = 3 + row_offset
            last_row_idx  = 3 + row_offset + n_day - 1
            merged_conf = tbl.rows[first_row_idx].cells[22].merge(
                          tbl.rows[last_row_idx].cells[22])
            sbg(merged_conf,"E8F5E9")
            sbo(merged_conf)
            smg(merged_conf,12,12,18,18)
            merged_conf.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            add_confidence_dropdown(merged_conf)

        row_offset += n_day

    # Col 0 en-tête : fond blanc sans couleur
    for row_i in range(2):
        c = tbl.rows[row_i].cells[0]
        sbg(c,"FFFFFF"); sbo(c)

def _safe(lst):
    """Remplace None/NaN par 0 pour matplotlib."""
    return [v if (v is not None and not (isinstance(v, float) and math.isnan(v))) else 0 for v in lst]

def gen_chart(df_sorted):
    import matplotlib; matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.ticker as ticker
    import numpy as np

    times = pd.to_datetime(df_sorted["valid_local"])
    speed = _safe(df_sorted["wind10_spd_kt"].tolist())
    gust  = _safe(df_sorted["wind10_gust_kt"].tolist())
    mslp  = _safe(df_sorted["mslp_hpa"].tolist())
    x     = list(range(len(times)))

    # ── Labels axe X : style du graphique de référence ───────────────────
    # Premier pas de chaque jour → "Day. DD\nMon\nHH:00"
    # Autres pas → "HH:00"
    xtick_labels = []
    prev_day = None
    for i, t in enumerate(times):
        if t.date() != prev_day:
            xtick_labels.append(
                f"{ENG_DS[t.weekday()]} {t.day}\n{ENG_MS[t.month]}\n{t.strftime('%H:%M')}"
            )
            prev_day = t.date()
        else:
            xtick_labels.append(t.strftime("%H:%M"))

    # ── Figure ────────────────────────────────────────────────────────────
    fig, ax1 = plt.subplots(figsize=(18, 6))
    fig.patch.set_facecolor("white")
    ax1.set_facecolor("white")

    # ── Axe droit : MSLP barres jaunes ───────────────────────────────────
    ax2 = ax1.twinx()
    ax2.bar(x, mslp, color="#FFC000", alpha=0.85, label="MSLP", zorder=2, width=0.7)
    ax2.set_ylabel("Pressure(hPa)", fontsize=14, fontweight="bold", color="#000000", labelpad=8)
    ax2.tick_params(axis="y", labelsize=13, colors="#000000")
    mc = [v for v in mslp if v]
    if mc:
        ax2.set_ylim(min(mc) - 3, max(mc) + 3)
    ax2.yaxis.set_major_formatter(ticker.FormatStrFormatter('%d'))
    ax2.spines["right"].set_color("#AAAAAA")

    # ── Axe gauche : courbes vent ─────────────────────────────────────────
    ax1.set_zorder(ax2.get_zorder() + 1)
    ax1.patch.set_visible(False)
    ax1.plot(x, speed, color="#4472C4", linewidth=2.0, label="Speed", zorder=5)
    ax1.plot(x, gust,  color="#ED7D31", linewidth=2.0, label="Gust",  zorder=5)
    ax1.set_ylabel("Wind (Kts)", fontsize=14, fontweight="bold", color="#000000", labelpad=8)
    ax1.tick_params(axis="y", labelsize=13, colors="#000000")
    gc = [v for v in gust if v]
    ax1.set_ylim(0, (max(gc) if gc else 20) + 4)
    ax1.set_xlim(-0.5, len(x) - 0.5)
    ax1.grid(True, linestyle="--", alpha=0.4, color="#CCCCCC", zorder=1)
    ax1.yaxis.set_major_formatter(ticker.FormatStrFormatter('%d'))
    ax1.spines["left"].set_color("#AAAAAA")
    ax1.spines["top"].set_visible(False)
    ax2.spines["top"].set_visible(False)

    # ── Axe X ─────────────────────────────────────────────────────────────
    ax1.set_xticks(x)
    ax1.set_xticklabels(xtick_labels, fontsize=11, color="#000000",
                        ha="center", va="top", multialignment="center")
    ax1.tick_params(axis="x", which="major", length=4, color="#AAAAAA", pad=4)
    ax1.set_xlabel("Times", fontsize=14, fontweight="bold", labelpad=10)

    # ── Titre ─────────────────────────────────────────────────────────────
    ax1.set_title("Wind at 10m and MSLP", fontsize=16, fontweight="bold", pad=12)

    # ── Légende centrée en bas ────────────────────────────────────────────
    l1, lb1 = ax1.get_legend_handles_labels()
    l2, lb2 = ax2.get_legend_handles_labels()
    ax1.legend(l2 + l1, lb2 + lb1,
               loc="lower center", ncol=3, fontsize=13,
               frameon=True, framealpha=0.95, edgecolor="#CCCCCC",
               bbox_to_anchor=(0.5, -0.38))

    plt.tight_layout(pad=1.2)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf
def sec_land(doc):
    s=doc.add_section(); s.orientation=WD_ORIENT.LANDSCAPE; s.page_width=Cm(29.7); s.page_height=Cm(21); s.left_margin=Cm(1.2); s.right_margin=Cm(1.2); s.top_margin=Cm(1.2); s.bottom_margin=Cm(1.2); return s

def sec_port(doc):
    s=doc.add_section(); s.orientation=WD_ORIENT.PORTRAIT; s.page_width=Cm(21); s.page_height=Cm(29.7); s.left_margin=Cm(1.5); s.right_margin=Cm(1.5); s.top_margin=Cm(1.5); s.bottom_margin=Cm(1.5); return s

# =============================================================================
# TEXTES DYNAMIQUES — Met Situation & Weather
# =============================================================================

def generate_met_situation(df: pd.DataFrame, run_datetime: datetime) -> str:
    """
    Génère automatiquement le texte Met Situation.
    Analyse : MSLP (anticyclone Sainte-Hélène), vent dominant (FIT/mousson).
    """
    df = df.copy()
    df["_dt"] = pd.to_datetime(df["valid_local"])

    mslp_vals  = df["mslp_hpa"].dropna()
    mslp_mean  = mslp_vals.mean()
    mslp_min   = mslp_vals.min()
    mslp_max   = mslp_vals.max()

    wind_vals  = df["wind10_spd_kt"].dropna()
    wind_mean  = wind_vals.mean()

    # Direction dominante sur la période
    dir_counts = df["wind10_dir"].dropna().value_counts()
    dom_dir    = dir_counts.index[0] if not dir_counts.empty else "—"

    # Détermination du régime synoptique
    # Anticyclone Sainte-Hélène actif si MSLP > 1013 hPa en moyenne
    if mslp_mean >= 1013:
        anticyclone = (f"The Saint Helena Anticyclone is active over the South Atlantic, "
                       f"maintaining mean sea level pressure around {mslp_mean:.0f} hPa over the area.")
    else:
        anticyclone = (f"The Saint Helena Anticyclone influence is weakened, "
                       f"with mean sea level pressure averaging {mslp_mean:.0f} hPa over the forecast period.")

    # FIT / Mousson : vent dominant du SW ou S indique mousson active
    monsoon_dirs = {"SW","SSW","WSW","S","SSE","SE"}
    if dom_dir in monsoon_dirs and wind_mean >= 8:
        fit_monsoon = (f"The Inter-Tropical Front (ITF) is positioned north of the area. "
                       f"The West African Monsoon is active, with dominant {dom_dir} winds "
                       f"averaging {wind_mean:.0f} knots, bringing moist maritime air onshore.")
    elif dom_dir in monsoon_dirs:
        fit_monsoon = (f"The West African Monsoon flow is present but moderate, "
                       f"with {dom_dir} winds averaging {wind_mean:.0f} knots. "
                       f"The ITF remains in the background of the forecast area.")
    else:
        fit_monsoon = (f"Dominant winds from {dom_dir} suggest limited monsoon penetration. "
                       f"The ITF position may be south of the area, "
                       f"with mean wind speed of {wind_mean:.0f} knots.")

    return f"{anticyclone} {fit_monsoon}"


def generate_weather_summary(df: pd.DataFrame, run_datetime: datetime) -> str:
    """
    Génère automatiquement le texte Weather — un paragraphe naturel,
    style bulletin météo, sans chiffres ni bullets.
    """
    df = df.copy()
    df["_dt"] = pd.to_datetime(df["valid_local"])
    df["_day"] = df["_dt"].dt.date
    days = sorted(df["_day"].unique())

    def sky_desc(rain_max):
        if rain_max is None or (isinstance(rain_max, float) and math.isnan(rain_max)):
            return "partly cloudy"
        if rain_max >= 70: return "overcast with rain expected"
        if rain_max >= 50: return "cloudy with possible showers"
        if rain_max >= 30: return "partly cloudy with isolated showers"
        return "mostly sunny to partly cloudy"

    def wind_desc(wind_max):
        if wind_max is None or (isinstance(wind_max, float) and math.isnan(wind_max)):
            return "light winds"
        if wind_max >= 20: return "strong winds"
        if wind_max >= 15: return "moderate to fresh winds"
        if wind_max >= 10: return "light to moderate winds"
        return "light winds"

    def sea_desc(swh_max):
        if swh_max is None or (isinstance(swh_max, float) and math.isnan(swh_max)):
            return "slight to moderate seas"
        if swh_max >= 2.0: return "rough seas"
        if swh_max >= 1.6: return "moderate to rough seas"
        if swh_max >= 1.0: return "moderate seas"
        return "slight seas"

    def risk_desc(swh_max, wind_max):
        if swh_max is None or (isinstance(swh_max, float) and math.isnan(swh_max)):
            return ""
        if swh_max >= 2.0 or (wind_max and not math.isnan(float(wind_max)) and wind_max >= 20):
            return " Mariners are advised to exercise extreme caution."
        if swh_max >= 1.6 or (wind_max and not math.isnan(float(wind_max)) and wind_max >= 15):
            return " Caution is advised for small crafts."
        return ""

    sentences = []
    for i, day in enumerate(days):
        day_df   = df[df["_day"] == day]
        dt_obj   = pd.Timestamp(day)
        day_lbl  = ENG_DF[dt_obj.weekday()]

        swh_max  = day_df["swh_m"].dropna().max()         if "swh_m"         in day_df.columns else None
        wind_max = day_df["wind10_spd_kt"].dropna().max()  if "wind10_spd_kt" in day_df.columns else None
        rain_max = day_df["rain_pct"].dropna().max()       if "rain_pct"      in day_df.columns else None

        sky  = sky_desc(rain_max)
        wind = wind_desc(wind_max)
        sea  = sea_desc(swh_max)
        risk = risk_desc(swh_max, wind_max)

        if i == 0:
            prefix = "Tonight into"
        elif i == 1:
            prefix = "On"
        else:
            prefix = "By"

        sentences.append(
            f"{prefix} {day_lbl}, {sky} skies are expected with {wind} and {sea}.{risk}"
        )

    return " ".join(sentences)


def generate_word_bulletin(df, run_datetime, warning_text=None, output_path=None):
    # Nom du fichier avec date et run — identique à l'Excel
    date_str = run_datetime.strftime("%d%m%Y")
    run_str  = f"{run_datetime.hour:02d}Z"
    if output_path is None:
        output_path = f"D:/pipeline/Marine_forecast_{date_str}_{run_str}.docx"
    else:
        base = output_path.replace(".docx", "").replace(".xlsx", "")
        output_path = f"{base}_{date_str}_{run_str}.docx"

    df_sorted = df.sort_values("valid_local").reset_index(drop=True)
    dt_start  = pd.to_datetime(df_sorted["valid_local"].iloc[0])
    dt_end    = dt_start + timedelta(days=3)

    doc = Document()
    # Section 1 — Paysage
    s1=doc.sections[0]; s1.orientation=WD_ORIENT.LANDSCAPE; s1.page_width=Cm(29.7); s1.page_height=Cm(21); s1.left_margin=Cm(1.2); s1.right_margin=Cm(1.2); s1.top_margin=Cm(1.2); s1.bottom_margin=Cm(1.2)

    # En-tête
    add_header(doc); ybar(doc)

    # Période — 12pt
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(2)
    ar(p,f"From {fmt_period(dt_start)} to {fmt_period(dt_end)} (local time)",sz=12,bold=True)
    ybar(doc)

    # Warning — 12pt
    pw=doc.add_paragraph(); pw.paragraph_format.space_before=Pt(3); pw.paragraph_format.space_after=Pt(3)
    pp=pw._p.get_or_add_pPr(); wd=warning_text or "Warning: None."
    bg="FF0000" if "ALERT" in wd.upper() else ("E2EFDA" if "Warning: None" in wd else "FFCC00")
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),bg); pp.append(s)
    r=pw.add_run(wd); r.font.size=Pt(12); r.font.bold=True
    if bg=="FF0000": r.font.color.rgb=WHITE
    elif bg=="E2EFDA": r.font.color.rgb=RGBColor(0x1E,0x6B,0x3C)
    else: r.font.color.rgb=RGBColor(0x7F,0x4E,0x00)

    # Met Situation — dynamique
    print("  📝 Génération Met Situation...")
    met_sit_text = generate_met_situation(df_sorted, run_datetime)
    pm=doc.add_paragraph(); pm.paragraph_format.space_before=Pt(3); pm.paragraph_format.space_after=Pt(1)
    ar(pm,"Met Situation:",sz=12,bold=True)
    pm2=doc.add_paragraph(); pm2.paragraph_format.space_before=Pt(0); pm2.paragraph_format.space_after=Pt(2)
    r_m=pm2.add_run(met_sit_text); r_m.font.size=Pt(11)

    # Weather — dynamique, paragraphe naturel
    print("  📝 Génération Weather summary...")
    weather_text = generate_weather_summary(df_sorted, run_datetime)
    pw2=doc.add_paragraph(); pw2.paragraph_format.space_before=Pt(3); pw2.paragraph_format.space_after=Pt(1)
    ar(pw2,"Weather:",sz=12,bold=True)
    pw3=doc.add_paragraph(); pw3.paragraph_format.space_before=Pt(0); pw3.paragraph_format.space_after=Pt(2)
    r_w=pw3.add_run(weather_text); r_w.font.size=Pt(11)

    # Titre Table 1 — 12pt
    pt=doc.add_paragraph(); pt.paragraph_format.space_before=Pt(4); pt.paragraph_format.space_after=Pt(2)
    ar(pt,"Table 1: Weather and Ocean parameters",sz=12,bold=True)

    # Tableau complet
    make_full_table(doc,df_sorted)

    # Légende
    pl=doc.add_paragraph(); pl.paragraph_format.space_before=Pt(2)
    ar(pl,"🟨 Weather condition: select from dropdown  |  🟩 Confidence: enter manually  |  Dir=Direction | Spd=Speed (kts for wind, m/s for currents) | Per=Period(s) | Ht=Height(m) | S.W=Significant Wave Height",sz=8,italic=True,color=RGBColor(0x66,0x66,0x66))

    # ── Page portrait : Figure 1 + Marées ────────────────────────────────
    sec_port(doc)
    print("  🖼️  Figure 1...")
    buf=gen_chart(df_sorted)
    pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; pc.paragraph_format.space_before=Pt(4); pc.paragraph_format.space_after=Pt(2)
    pc.add_run().add_picture(buf,width=Cm(17))
    # Titre sous le graphique — 12pt
    pf1=doc.add_paragraph(); pf1.alignment=WD_ALIGN_PARAGRAPH.CENTER; pf1.paragraph_format.space_before=Pt(2); pf1.paragraph_format.space_after=Pt(6)
    ar(pf1,f"Figure 1: Trends of winds and Mean Sea Level Pressure from {ordinal(dt_start.day)} to {ordinal(dt_end.day)} {ENG_MF[dt_start.month]} {dt_start.year}",sz=12,bold=True)

    # Titre Table 2 — 12pt
    pt2=doc.add_paragraph(); pt2.paragraph_format.space_before=Pt(4); pt2.paragraph_format.space_after=Pt(2)
    ar(pt2,"❖  Table 2: Tides predictions for SEME",sz=12,bold=True)

    # ── Récupérer les marées via WorldTides API ───────────────────────────
    try:
        import tides as tides_module
        tide_data  = tides_module.get_tides(dt_start, days=4)
        tide_rows  = tides_module.format_tide_table(tide_data, dt_start, days=4)
    except Exception as e:
        print(f"  ⚠️  Marées non disponibles : {e}")
        tide_rows = []

    n_data = max(len(tide_rows), 4)
    tide=doc.add_table(rows=2+n_data, cols=9)
    tide.style='Table Grid'; tide.alignment=WD_TABLE_ALIGNMENT.CENTER

    # En-tête ligne 1
    r0=tide.rows[0]
    sbg(r0.cells[0],"FFFFFF"); sbo(r0.cells[0]); cp(r0.cells[0],"Dates",sz=9,bold=True,color=RGBColor(0,0,0))
    ch=r0.cells[1].merge(r0.cells[4]); sbg(ch,"FFFFFF"); sbo(ch); cp(ch,"HIGH TIDES",sz=9,bold=True,color=RGBColor(0,0,0))
    cl2=r0.cells[5].merge(r0.cells[8]); sbg(cl2,"FFFFFF"); sbo(cl2); cp(cl2,"LOW TIDES",sz=9,bold=True,color=RGBColor(0,0,0))

    # En-tête ligne 2
    r1=tide.rows[1]
    for j,lbl in enumerate(["","Time","Height","Time","Height","Time","Height","Time","Height"]):
        c=r1.cells[j]; sbg(c,"FFFFFF"); sbo(c)
        cp(c,lbl,sz=9,bold=True,color=RGBColor(0,0,0))

    # Données : remplies automatiquement si WorldTides disponible, sinon vides
    DAY_BGS = ["E2EFDA","FFFFFF","E2EFDA","FFFFFF"]
    for i in range(n_data):
        row=tide.rows[2+i]; bg=DAY_BGS[i%4]
        if i < len(tide_rows):
            tr = tide_rows[i]
            vals=[tr["date"],
                  tr["high1"]["time"],tr["high1"]["height"],
                  tr["high2"]["time"],tr["high2"]["height"],
                  tr["low1"]["time"], tr["low1"]["height"],
                  tr["low2"]["time"], tr["low2"]["height"]]
        else:
            vals=[""] * 9
        for j,v in enumerate(vals):
            c=row.cells[j]; sbg(c,bg); sbo(c); cp(c,str(v),sz=9)

    # ── Page portrait : Figure 2 ENSgram + commentaire côte à côte ──────────
    sec_port(doc)
    print("  🌐 Téléchargement Figure 2 (ENSgram ECMWF)...")
    ecmwf_urls  = build_ecmwf_urls(run_datetime)
    img_ensgram = fetch_ecmwf_image(ecmwf_urls["ensgram"], "ENSgram")

    # Tableau 2 colonnes : image gauche (12cm) | commentaire droite (5.5cm)
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_img  = tbl2.rows[0].cells[0]
    cell_com  = tbl2.rows[0].cells[1]

    # Supprimer bordures visibles
    for c in [cell_img, cell_com]:
        snb(c)

    # ── Colonne gauche : ENSgram ──
    cell_img.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if cell_img.paragraphs:
        cell_img.paragraphs[0]._p.getparent().remove(cell_img.paragraphs[0]._p)

    if img_ensgram:
        pi2 = cell_img.add_paragraph()
        pi2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi2.paragraph_format.space_before = Pt(2)
        pi2.paragraph_format.space_after  = Pt(2)
        pi2.add_run().add_picture(img_ensgram, width=Cm(11.5))
    else:
        pi2 = cell_img.add_paragraph()
        pi2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ph = pi2.add_run("[ ENSgram — non disponible ]")
        r_ph.font.size = Pt(9); r_ph.font.italic = True; r_ph.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

    # Légende sous l'image
    pleg = cell_img.add_paragraph()
    pleg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pleg.paragraph_format.space_before = Pt(2)
    r_leg = pleg.add_run(f"Figure 2: Wave ensemblegram issued by ECMWF from {ordinal(dt_start.day)} to {ordinal(dt_end.day)} {ENG_MF[dt_start.month]} {dt_start.year}")
    r_leg.font.size = Pt(9); r_leg.font.bold = True

    # ── Colonne droite : commentaire automatique tendance SWH ──
    cell_com.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if cell_com.paragraphs:
        cell_com.paragraphs[0]._p.getparent().remove(cell_com.paragraphs[0]._p)

    # Calcul automatique tendance SWH
    swh_vals   = df_sorted["swh_m"].dropna()
    swh_max    = swh_vals.max()
    swh_min    = swh_vals.min()
    idx_max    = swh_vals.idxmax()
    peak_dt    = pd.to_datetime(df_sorted.loc[idx_max, "valid_local"])
    n          = len(swh_vals)
    first_half = swh_vals.iloc[:n//2].mean()
    second_half= swh_vals.iloc[n//2:].mean()
    if second_half > first_half + 0.1:
        trend_word = "increasing"
        trend_arrow = "↗"
    elif second_half < first_half - 0.1:
        trend_word = "decreasing"
        trend_arrow = "↘"
    else:
        trend_word = "steady"
        trend_arrow = "→"

    ensgram_comment = (
        f"SWH {trend_arrow} {trend_word}. "
        f"Peak: {swh_max:.1f} m ({peak_dt.strftime('%a %d %b, %HZ')}). "
        f"Min: {swh_min:.1f} m."
    )

    # Titre petit box
    pt_com = cell_com.add_paragraph()
    pt_com.paragraph_format.space_before = Pt(4)
    pt_com.paragraph_format.space_after  = Pt(3)
    r_tc = pt_com.add_run("❖  ENSgram summary:")
    r_tc.font.size = Pt(9); r_tc.font.bold = True

    # Box coloré selon tendance
    pc_com = cell_com.add_paragraph()
    pc_com.paragraph_format.space_before = Pt(2)
    pc_com.paragraph_format.space_after  = Pt(2)
    r_com  = pc_com.add_run(ensgram_comment)
    r_com.font.size = Pt(10); r_com.font.bold = True

    # ── Page portrait : Figure 3 (4 cartes ECMWF) ────────────────────────
    sec_port(doc)
    fig3d=run_datetime+timedelta(days=1)
    print("  🌐 Téléchargement Figure 3 (4 cartes ECMWF)...")

    FIG3_KEYS   = ["swh_dir", "swh_prob", "mwp_8s", "mwp_15s"]
    FIG3_LABELS = [
        "SWH & mean direction",
        "Prob. SWH ≥2m",
        "Prob. mean wave period ≥8s",
        "Prob. mean wave period ≥15s",
    ]
    fig3_imgs = [fetch_ecmwf_image(ecmwf_urls[k], FIG3_LABELS[i])
                 for i, k in enumerate(FIG3_KEYS)]

    it=doc.add_table(rows=2,cols=2); it.style='Table Grid'; it.alignment=WD_TABLE_ALIGNMENT.CENTER
    cells=[it.rows[0].cells[0],it.rows[0].cells[1],it.rows[1].cells[0],it.rows[1].cells[1]]

    for idx, cell in enumerate(cells):
        snb(cell); cell.width=Cm(8.5)
        # Supprimer paragraphe vide initial
        if cell.paragraphs:
            cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)

        # Sous-titre de la carte
        pt_sub = cell.add_paragraph()
        pt_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt_sub.paragraph_format.space_before = Pt(2)
        pt_sub.paragraph_format.space_after  = Pt(1)
        r_sub = pt_sub.add_run(FIG3_LABELS[idx])
        r_sub.font.size  = Pt(7)
        r_sub.font.bold  = True

        # Image ou placeholder
        pi = cell.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.paragraph_format.space_before = Pt(1)
        pi.paragraph_format.space_after  = Pt(2)

        if fig3_imgs[idx]:
            pi.add_run().add_picture(fig3_imgs[idx], width=Cm(8.0))
        else:
            pp2 = pi._p.get_or_add_pPr()
            s2  = OxmlElement('w:shd'); s2.set(qn('w:val'),'clear'); s2.set(qn('w:color'),'auto'); s2.set(qn('w:fill'),'F0F8FF'); pp2.append(s2)
            sp2 = OxmlElement('w:spacing'); sp2.set(qn('w:before'),'800'); sp2.set(qn('w:after'),'800'); pp2.append(sp2)
            ri  = pi.add_run(f"[ Image {idx+1} — non disponible ]")
            ri.font.size = Pt(8); ri.font.italic = True; ri.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

    # Titre Figure 3 — sous le tableau d'images
    pf3=doc.add_paragraph(); pf3.alignment=WD_ALIGN_PARAGRAPH.CENTER; pf3.paragraph_format.space_before=Pt(4); pf3.paragraph_format.space_after=Pt(2)
    ar(pf3,f"Figure 3: Probability forecast for {ENG_DF[fig3d.weekday()]}, {ordinal(fig3d.day)} {ENG_MF[fig3d.month]} {fig3d.year} at 01:00 pm local time.",sz=12,bold=True)

    doc.save(output_path)
    print(f"  ✅ Bulletin Word : {output_path}")
    return output_path